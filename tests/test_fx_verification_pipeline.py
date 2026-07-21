# -*- coding: utf-8 -*-
"""
fx_verification_pipeline.py 테스트 스위트.

원칙:
- 네트워크 호출(수출입은행 API, Claude API)은 절대 실제로 하지 않는다.
  fetch_rates_for_date/anthropic.Anthropic을 monkeypatch로 대체한다.
- 순수 로직 함수는 모킹 없이 직접 검증한다.
- 새 픽스처를 만들기보다 sample_data/, robustness_test/의 기존 파일을 재사용한다.
"""
import os
import math
from pathlib import Path

import pandas as pd
import pytest

import fx_verification_pipeline as fxp

ROOT_DIR = Path(__file__).resolve().parent.parent
SAMPLE_DIR = ROOT_DIR / "sample_data"
ROBUST_DIR = ROOT_DIR / "robustness_test"


# ------------------------------------------------------------------
# 순수 로직 함수 - 모킹 불필요
# ------------------------------------------------------------------

class TestNormalizeRate:
    def test_jpy_divides_by_100(self):
        assert fxp.normalize_rate(937.73, "JPY") == pytest.approx(9.3773)

    def test_usd_passthrough(self):
        assert fxp.normalize_rate(1430.20, "USD") == 1430.20


class TestNormalizeAccountCode:
    @pytest.mark.parametrize("raw, expected", [
        ("108", "108"),
        (108, "108"),
        (108.0, "108"),
        ("108.0", "108"),
        (" 108 ", "108"),
        ("ABC", "ABC"),  # 숫자로 변환 불가능한 코드는 원본을 그대로 유지
    ])
    def test_various_formats(self, raw, expected):
        assert fxp._normalize_account_code(raw) == expected


class TestResolveCurUnit:
    def test_known_currency(self):
        cur_unit, warning = fxp.resolve_cur_unit("USD")
        assert cur_unit == "USD"
        assert warning is None

    def test_unknown_currency_returns_warning(self):
        cur_unit, warning = fxp.resolve_cur_unit("XYZ")
        assert cur_unit == "XYZ"
        assert "매핑 미등록" in warning


class TestCandidateWrongDates:
    def test_month_boundary(self):
        # 결제일이 1월 3일이면 '전월말'은 전년도 12월 31일이어야 한다.
        candidates = fxp._candidate_wrong_dates(pd.Timestamp("2025-01-03"))
        candidate_strs = {d.strftime("%Y-%m-%d") for d in candidates}
        assert "2024-12-31" in candidate_strs

    def test_all_candidates_before_settle_date(self):
        settle_date = pd.Timestamp("2025-05-15")
        candidates = fxp._candidate_wrong_dates(settle_date)
        assert all(d < settle_date for d in candidates)
        assert len(candidates) == len(set(candidates))  # 중복 없음


class TestCheckAggregateMateriality:
    def _screened(self, diffs):
        return pd.DataFrame({"회사계상액과의차이(KRW)": diffs})

    def test_within_tolerance(self):
        agg = fxp.check_aggregate_materiality(self._screened([1000000, -900000, 999999]), ampt=3000000)
        assert not agg["중요성초과여부"]
        assert agg["확인불가건수"] == 0

    def test_breach_on_gross_sum(self):
        # 순합계는 작아도(방향이 반대라 상쇄) 절대값 합계 기준으로는 AMPT를 넘는 경우
        agg = fxp.check_aggregate_materiality(self._screened([2000000, -2000000, 2500000]), ampt=3000000)
        assert agg["중요성초과여부"]

    def test_none_values_excluded_but_counted(self):
        # 환율조회 실패로 None이 된 행은 합계에서 빠지되, 확인불가건수로 드러나야 한다.
        agg = fxp.check_aggregate_materiality(self._screened([1000000, None, None]), ampt=3000000)
        assert agg["순차이합계(KRW)"] == 1000000
        assert agg["확인불가건수"] == 2


class TestRowStatus:
    def test_flag_keyword(self):
        assert fxp._row_status(["이상치(정밀검증필요)"]) == "flag"

    def test_ok_keyword(self):
        assert fxp._row_status(["적정(스크리닝통과)"]) == "ok"

    def test_neutral(self):
        assert fxp._row_status(["그 외 아무 값"]) == "neutral"

    def test_outlier_count_header_flags_even_without_keyword(self):
        assert fxp._row_status([3], headers=["이상치건수"]) == "flag"
        assert fxp._row_status([0], headers=["이상치건수"]) == "neutral"


class TestAutoAssignVoucherNumber:
    def test_groups_balanced_two_line_vouchers(self):
        df = pd.DataFrame({
            "차변금액": [100000, None, 50000, None],
            "대변금액": [None, 100000, None, 50000],
        })
        result = fxp.auto_assign_voucher_number(df)
        assert list(result["전표번호(거래ID)"]) == ["AUTO00001", "AUTO00001", "AUTO00002", "AUTO00002"]

    def test_raises_when_last_voucher_unbalanced(self):
        df = pd.DataFrame({"차변금액": [100000, None, 50000], "대변금액": [None, 100000, None]})
        with pytest.raises(ValueError):
            fxp.auto_assign_voucher_number(df)


class TestDemoAccountMappingAndVoucherNumberingScript:
    """scripts/demo_account_mapping_and_voucher_numbering.py는 build_account_mapping/
    apply_account_mapping/auto_assign_voucher_number - main()/CLI에 안 붙어있어 --workpaper
    한 번으로는 못 보여주는 두 '수동 개입' 기능 - 을 실제로 어떻게 호출하는지 보여주는
    데모다. 콘솔 출력이 아니라 함수가 실제로 반환하는 값을 assert한다."""

    @staticmethod
    def _load_demo_module():
        import importlib.util
        script_path = ROOT_DIR / "scripts" / "demo_account_mapping_and_voucher_numbering.py"
        spec = importlib.util.spec_from_file_location("demo_account_mapping_and_voucher_numbering", script_path)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        return module

    def test_account_mapping_demo_auto_confirms_and_flags_ambiguous(self):
        demo = self._load_demo_module()
        mapping_result, confirmed_mapping, mapped = demo.demo_account_mapping()

        by_code = mapping_result.set_index("회사계정코드")
        assert by_code.loc["1081", "상태"] == "자동확정"
        assert by_code.loc["1081", "추천표준분류"] == "108"
        assert by_code.loc["9999", "상태"] == "확인필요(매칭 없음)"

        # 사람이 "9999"(달러채권)를 108로 확정했다고 가정한 뒤 적용한 결과가 실제로 반영됐는지.
        assert confirmed_mapping["9999"] == "108"
        mapped_by_orig = mapped.set_index("원본계정코드")
        assert mapped_by_orig.loc["9999", "계정코드"] == "108"
        assert mapped_by_orig.loc["9999", "계정과목"] == "외상매출금(외화)"

    def test_voucher_numbering_demo_assigns_balanced_groups(self):
        demo = self._load_demo_module()
        result = demo.demo_auto_voucher_numbering()
        assert list(result["전표번호(거래ID)"]) == [
            "AUTO00001", "AUTO00001", "AUTO00002", "AUTO00002",
        ]


# ------------------------------------------------------------------
# 지저분한 원본 파일 대응 (robustness_test/messy_*.xlsx 재사용)
# ------------------------------------------------------------------

class TestLoadJournalRobustness:
    def test_messy_journal_loads_with_header_detection_and_normalization(self):
        path = ROBUST_DIR / "messy_분개장.xlsx"
        if not path.exists():
            pytest.skip("robustness_test/messy_분개장.xlsx 없음")
        df = fxp.load_journal(str(path))
        # 제목행/빈행이 섞여 있어도 필수 컬럼이 정상적으로 인식되어야 한다.
        for col in fxp.REQUIRED_JOURNAL_COLUMNS:
            assert col in df.columns
        # 계정코드는 전부 문자열로 정규화되어야 한다 ("108.0" 같은 형태가 남아있으면 안 됨).
        assert df["계정코드"].apply(lambda v: not str(v).endswith(".0")).all()


class TestMessySampleEndToEnd:
    """messy_분개장.xlsx뿐 아니라 messy_명세서_외화자산부채명세서.xlsx/messy_계정별원장.xlsx까지
    묶어서, 지저분한 3파일 세트 전체가 파이프라인을 끝까지 죽지 않고 통과하는지 확인한다.
    (messy_분개장.xlsx만 헤더 자동인식 대상이고, 명세서/원장은 1행이 곧바로 헤더라 지금
    로더로도 바로 읽힌다 - generate_messy_sample.py 확인 결과.) TXN101/105 조합으로 재현된
    전표번호 중복 오채번의 데이터 무결성 경고도 함께 확인한다."""

    @pytest.fixture
    def messy_inputs(self):
        journal_path = ROBUST_DIR / "messy_분개장.xlsx"
        schedule_path = ROBUST_DIR / "messy_명세서_외화자산부채명세서.xlsx"
        ledger_path = ROBUST_DIR / "messy_계정별원장.xlsx"
        if not (journal_path.exists() and schedule_path.exists() and ledger_path.exists()):
            pytest.skip("robustness_test/messy_*.xlsx 3종 세트가 없음")
        journal = fxp.load_journal(str(journal_path))
        schedule = pd.read_excel(schedule_path)
        ledger = pd.read_excel(ledger_path)
        return journal, schedule, ledger

    def test_full_pipeline_survives_messy_three_file_set(self, messy_inputs, monkeypatch):
        journal, schedule, ledger = messy_inputs

        # messy 샘플 자체의 환율(TXN101~109)만 있으면 되므로, 못 찾는 날짜는 그냥 빈 값으로
        # 둬서 "확인불가/오류"로 우아하게 처리되는지도 같이 확인한다(전부 맞출 필요 없음).
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch({}))

        settlements = fxp.extract_settlement_transactions(journal)
        screened = fxp.screen_fx_settlements(settlements)
        ref_check = fxp.detect_reference_date_mismatch(settlements)
        agg = fxp.check_aggregate_materiality(screened, ampt=3000000)

        ye_txns = fxp.extract_yearend_transactions(journal)
        missing_ye = fxp.get_unsettled_yearend_candidates(journal, schedule)
        ye_result = fxp.verify_yearend_translation(ye_txns, missing_ye, "2025-12-31")

        ledger_result = fxp.verify_ledger_reconciliation(journal, ledger)
        rollforward_result = fxp.build_rollforward_verification(journal, schedule, screened)

        # 전체 체인이 예외 없이 끝까지 돌고, 각 단계가 빈 결과를 내지 않아야 한다.
        assert not settlements.empty
        assert not screened.empty
        assert not ref_check.empty
        assert agg is not None
        assert not ledger_result.empty
        assert not rollforward_result.empty

        # TXN101 발생 라인에 EUR/Munich Parts AG 거래가 오채번으로 섞여 있으므로
        # (통화·거래처가 다른 발생 라인 2건 이상) 데이터 무결성 경고가 떠야 한다.
        warning = settlements.loc[settlements["거래ID"] == "TXN101", "데이터무결성경고"]
        assert warning.notna().any()
        assert "전표번호 중복 의심" in warning.dropna().iloc[0]

    def test_workpaper_generation_survives_messy_input(self, messy_inputs, monkeypatch, tmp_path):
        """--workpaper까지 포함해 지저분한 입력으로도 감사조서 생성 자체가 죽지 않는지 확인."""
        journal, schedule, ledger = messy_inputs
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch({}))
        monkeypatch.setattr(fxp, "EVIDENCE_DIR", str(tmp_path / "no_evidence"))

        settlements = fxp.extract_settlement_transactions(journal)
        screened = fxp.screen_fx_settlements(settlements)
        ref_check = fxp.detect_reference_date_mismatch(settlements)
        agg = fxp.check_aggregate_materiality(screened, ampt=3000000)
        verified = fxp.verify_with_evidence(screened)

        ye_txns = fxp.extract_yearend_transactions(journal)
        missing_ye = fxp.get_unsettled_yearend_candidates(journal, schedule)
        ye_result = fxp.verify_yearend_translation(ye_txns, missing_ye, "2025-12-31")

        ledger_result = fxp.verify_ledger_reconciliation(journal, ledger)
        rollforward_result = fxp.build_rollforward_verification(journal, schedule, screened)

        output_path = tmp_path / "감사조서_messy_테스트.docx"
        fxp.generate_audit_workpaper(
            str(output_path),
            screened=screened, agg=agg, ref_check=ref_check, verified=verified,
            ye_result=ye_result, ledger_result=ledger_result,
            rollforward_result=rollforward_result, year_end_date="2025-12-31",
        )
        assert output_path.exists()


# ------------------------------------------------------------------
# fetch_rates_for_date monkeypatch로 테스트 - 네트워크 호출 없음
# ------------------------------------------------------------------

def _make_fake_fetch(rate_table_by_date: dict):
    def _fake_fetch(date_str, _retries=2):
        return rate_table_by_date.get(date_str, {})
    return _fake_fetch


def _settlement_row(**overrides):
    row = {
        "거래ID": "T1", "결제일": pd.Timestamp("2025-05-15"), "통화": "USD", "구분": "채권",
        "거래처": "테스트거래처", "데이터무결성경고": None,
        "외화금액": 20000, "결제원화금액": 20000 * 1438.50, "발생원화금액": 20000 * 1460.80,
        "회사적용환율(내재)": 1438.50, "회사계상_외환차익차손": 20000 * (1438.50 - 1460.80),
    }
    row.update(overrides)
    return row


class TestScreenFxSettlements:
    def test_flags_deviation_over_tolerance(self, monkeypatch):
        # 공식환율(1000)과 회사환율(1438.5)이 5%를 훨씬 넘게 벌어진 케이스
        monkeypatch.setattr(fxp, "fetch_rates_for_date",
                             _make_fake_fetch({"20250515": {"USD": 1000.0}}))
        settlements = pd.DataFrame([_settlement_row()])
        result = fxp.screen_fx_settlements(settlements)
        assert result.iloc[0]["1차플래그"] == "이상치(정밀검증필요)"

    def test_txn002_style_case_passes_individual_but_shows_diff(self, monkeypatch):
        # 실제 결제일 환율(1415.80) 대신 전월말 환율(1438.50)을 잘못 쓴 경우
        # 괴리율은 5% 미만이라 개별 기준은 통과해야 한다 (합계 중요성에서 걸리는 게 설계 의도).
        monkeypatch.setattr(fxp, "fetch_rates_for_date",
                             _make_fake_fetch({"20250515": {"USD": 1415.80}}))
        settlements = pd.DataFrame([_settlement_row()])
        result = fxp.screen_fx_settlements(settlements)
        assert result.iloc[0]["1차플래그"] == "적정(스크리닝통과)"
        assert result.iloc[0]["회사계상액과의차이(KRW)"] != 0

    def test_degrades_gracefully_when_rate_lookup_fails(self, monkeypatch):
        # API 키 누락/만료 등으로 환율을 아예 못 찾는 상황(5영업일 폴백까지 전부 실패)에도
        # 배치 전체가 죽지 않고 해당 행만 오류로 표시되어야 한다.
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch({}))
        settlements = pd.DataFrame([_settlement_row(), _settlement_row(거래ID="T2")])
        result = fxp.screen_fx_settlements(settlements)
        assert len(result) == 2
        assert (result["1차플래그"] == "오류(환율조회실패)").all()
        assert result["공식매매기준율"].isna().all()

    def test_unmapped_currency_skips_rate_lookup_entirely(self, monkeypatch):
        # CUR_UNIT_MAP에 없는 통화는 detect_reference_date_mismatch/verify_yearend_translation과
        # 동일하게, API가 우연히 받아줄 수 있어도 조회 자체를 하지 않고 확인불가로 남겨야 한다.
        def _fail_if_called(date_str, _retries=2):
            raise AssertionError("미등록 통화는 fetch_rates_for_date가 호출되면 안 된다")

        monkeypatch.setattr(fxp, "fetch_rates_for_date", _fail_if_called)
        settlements = pd.DataFrame([_settlement_row(통화="XYZ")])
        result = fxp.screen_fx_settlements(settlements)
        assert result.iloc[0]["1차플래그"].startswith("확인불가")
        assert result.iloc[0]["공식매매기준율"] is None
        assert result.iloc[0]["회사계상액과의차이(KRW)"] is None


class TestDetectReferenceDateMismatch:
    def test_finds_prior_month_end_misuse(self, monkeypatch):
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch({
            "20250515": {"USD": 1415.80},  # 실제 결제일 환율
            "20250430": {"USD": 1438.50},  # 회사가 실수로 사용한 전월말 환율
        }))
        settlements = pd.DataFrame([_settlement_row()])
        result = fxp.detect_reference_date_mismatch(settlements)
        row = result.iloc[0]
        assert row["기준일판정"].startswith("기준일 오류 의심")
        assert row["추정사용일자"] == "2025-04-30"

    def test_normal_case_confirms_settle_date_rate(self, monkeypatch):
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch({
            "20250515": {"USD": 1438.50},
        }))
        # 기본 _settlement_row()가 이미 회사적용환율(내재)=1438.50이므로 그대로 사용
        settlements = pd.DataFrame([_settlement_row()])
        result = fxp.detect_reference_date_mismatch(settlements)
        assert result.iloc[0]["기준일판정"] == "정상(결제일 환율 사용 확인)"

    def test_degrades_gracefully_when_rate_lookup_fails(self, monkeypatch):
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch({}))
        settlements = pd.DataFrame([_settlement_row()])
        result = fxp.detect_reference_date_mismatch(settlements)
        assert result.iloc[0]["기준일판정"].startswith("확인불가")


class TestVerifyWithEvidence:
    def test_no_evidence_file_marks_request_needed(self, tmp_path, monkeypatch):
        monkeypatch.setattr(fxp, "EVIDENCE_DIR", str(tmp_path))  # 빈 폴더
        flagged = pd.DataFrame([{"거래ID": "TXN999", "통화": "USD",
                                  "외화금액": 1000, "회사적용환율(내재)": 1400.0}])
        result = fxp.verify_with_evidence(flagged)
        assert result.iloc[0]["증빙상태"] == "증빙 요청 필요"

    def test_empty_input_returns_correctly_shaped_dataframe(self):
        # 이상치가 하나도 없는 정상적인 기간에도 다운스트림에서 컬럼 접근 시 에러가 나면 안 된다.
        empty = pd.DataFrame(columns=["거래ID", "통화", "외화금액", "회사적용환율(내재)"])
        result = fxp.verify_with_evidence(empty)
        assert result.empty
        for col in ("거래ID", "증빙상태", "증빙확인환율", "최종판정"):
            assert col in result.columns

    def test_ocr_call_failure_degrades_gracefully(self, tmp_path, monkeypatch):
        monkeypatch.setattr(fxp, "EVIDENCE_DIR", str(tmp_path))
        evidence_file = tmp_path / "TXN999.png"
        evidence_file.write_bytes(b"fake-image-bytes")

        def _raise(*args, **kwargs):
            raise RuntimeError("ANTHROPIC_API_KEY 미설정")

        monkeypatch.setattr(fxp, "extract_rate_from_evidence", _raise)
        flagged = pd.DataFrame([{"거래ID": "TXN999", "통화": "USD",
                                  "외화금액": 1000, "회사적용환율(내재)": 1400.0}])
        result = fxp.verify_with_evidence(flagged)
        assert result.iloc[0]["증빙상태"] == "OCR 호출 실패 - 수기 확인 필요"


class TestOcrRecheckSample:
    """'적정(증빙과 일치)' 판정 건 중 일부를 QC 표본으로 뽑는 _select_ocr_recheck_sample().
    고정 시드(OCR_RECHECK_SAMPLE_SEED)를 쓰므로 표본 크기는 결정론적으로 검증 가능하다."""

    def test_marks_ceil_of_rate_percent_of_matched_rows(self):
        rows = [{"거래ID": f"TXN{i}", "최종판정": "적정(증빙과 일치)"} for i in range(20)]
        fxp._select_ocr_recheck_sample(rows)
        sampled = [r for r in rows if r["OCR재확인표본"] is not None]
        expected = math.ceil(20 * fxp.OCR_RECHECK_SAMPLE_RATE)
        assert len(sampled) == expected

    def test_single_matched_row_still_gets_sampled(self):
        rows = [{"거래ID": "TXN1", "최종판정": "적정(증빙과 일치)"}]
        fxp._select_ocr_recheck_sample(rows)
        assert rows[0]["OCR재확인표본"] is not None

    def test_no_matched_rows_column_still_added_as_none(self):
        # 매칭된 행이 0건이어도 컬럼 자체는 항상 채워져야 build_fx_detail_report()의
        # 컬럼 선택(verified[[..., "OCR재확인표본"]])이 KeyError 없이 동작한다.
        rows = [{"거래ID": "TXN1", "최종판정": "부적정(증빙과 불일치 - 재계산 필요)"}]
        fxp._select_ocr_recheck_sample(rows)
        assert rows[0]["OCR재확인표본"] is None

    def test_non_matched_rows_are_never_sampled(self):
        rows = [
            {"거래ID": "TXN1", "최종판정": "적정(증빙과 일치)"},
            {"거래ID": "TXN2", "최종판정": "부적정(증빙과 불일치 - 재계산 필요)"},
            {"거래ID": "TXN3", "최종판정": "미확인(증빙 미확보)"},
        ]
        fxp._select_ocr_recheck_sample(rows)
        assert rows[1]["OCR재확인표본"] is None
        assert rows[2]["OCR재확인표본"] is None

    def test_sample_text_never_overlaps_flag_keywords(self):
        # Exception(빨간색 하이라이트) 목록에 QC 표본이 섞여 들어가면 안 되므로,
        # 표시 문구가 FLAG_KEYWORDS(이상치/부적정/오류/불일치/누락/미확인/재검토 필요/확인 필요)
        # 와 절대 겹치지 않아야 한다.
        rows = [{"거래ID": "TXN1", "최종판정": "적정(증빙과 일치)"}]
        fxp._select_ocr_recheck_sample(rows)
        text = rows[0]["OCR재확인표본"]
        assert not any(k in text for k in fxp.FLAG_KEYWORDS)


class TestOcrRecheckSampleIntegration:
    def test_ocr_recheck_column_flows_into_fx_detail_report(self, sample_pipeline_inputs,
                                                              monkeypatch, tmp_path):
        journal, _schedule = sample_pipeline_inputs
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch(SAMPLE_FAKE_RATES))
        # 증빙 폴더를 빈 임시경로로 돌려 OCR(실제 API) 호출 자체가 일어나지 않게 한다.
        monkeypatch.setattr(fxp, "EVIDENCE_DIR", str(tmp_path / "no_evidence"))

        settlements = fxp.extract_settlement_transactions(journal)
        screened = fxp.screen_fx_settlements(settlements)
        ref_check = fxp.detect_reference_date_mismatch(settlements)
        verified = fxp.verify_with_evidence(screened)
        fx_detail = fxp.build_fx_detail_report(screened, ref_check, verified)

        assert "OCR재확인표본" in verified.columns
        assert "OCR재확인표본" in fx_detail.columns


class _FakeTextBlock:
    type = "text"

    def __init__(self, text):
        self.text = text


class _FakeMessage:
    def __init__(self, text):
        self.content = [_FakeTextBlock(text)]


class _FakeMessages:
    def __init__(self, response_text):
        self._response_text = response_text

    def create(self, **kwargs):
        return _FakeMessage(self._response_text)


class _FakeAnthropicClient:
    def __init__(self, response_text):
        self.messages = _FakeMessages(response_text)


class TestExtractRateFromEvidenceParsing:
    """Claude Vision API 자체는 호출하지 않고, anthropic 클라이언트를 목킹해
    JSON 파싱 성공/실패 분기만 검증한다."""

    def _run_with_fake_response(self, monkeypatch, tmp_path, response_text):
        import anthropic
        img = tmp_path / "evidence.png"
        img.write_bytes(b"fake-image-bytes")
        monkeypatch.setattr(anthropic, "Anthropic",
                             lambda: _FakeAnthropicClient(response_text))
        return fxp.extract_rate_from_evidence(str(img))

    def test_parses_valid_json_response(self, monkeypatch, tmp_path):
        response = '{"거래일자": "2025-05-15", "통화": "USD", "적용환율": 1415.8, "외화금액": 20000, "원화금액": 28316000}'
        result = self._run_with_fake_response(monkeypatch, tmp_path, response)
        assert result["적용환율"] == 1415.8

    def test_unparseable_response_falls_back_to_error(self, monkeypatch, tmp_path):
        result = self._run_with_fake_response(monkeypatch, tmp_path, "이건 JSON이 아닙니다")
        assert "error" in result


# ------------------------------------------------------------------
# End-to-end 오라클 테스트: sample_data/ 14개 거래를 검증포인트_참고용.xlsx와 대조
# ------------------------------------------------------------------

# generate_samples.py에 기록된 실제/설계 매매기준율을 그대로 재현한 오프라인 환율표.
# 다른 후보일자(전월말/전영업일 등)는 의도적으로 비워둬 candidate 탐색이 '매칭 없음'으로
# 안전하게 종료되는지도 같이 검증한다.
SAMPLE_FAKE_RATES = {
    "20250425": {"USD": 1430.20},   # TXN001 결제일
    "20250515": {"USD": 1415.80},   # TXN002 실제 결제일 환율
    "20250430": {"USD": 1438.50},   # TXN002가 잘못 사용한 전월말 환율
    "20250718": {"JPY(100)": 937.73},  # TXN003 결제일
    "20250930": {"USD": 1402.20},   # TXN004 결제일
    "20250611": {"EUR": 1554.64},   # TXN005 결제일
    "20251231": {"USD": 1434.90, "HKD": 183.50, "SGD": 1065.90},  # 결산일 공식환율 (TXN006/008/010/014 판정 기준)
    "20250310": {"GBP": 1885.60},   # TXN009 분할결제 1차
    "20250520": {"GBP": 1901.30},   # TXN009 분할결제 2차
    "20250605": {"CHF": 1620.00},   # TXN011 결제일 - 회사 계상(1750.00)과 괴리 8%+ (우대환율)
    "20250702": {"CAD": 1008.40},   # TXN012 결제일
    "20250814": {"AUD": 918.75},    # TXN013 결제일 - 회사 계상(980.00)이 오류, 증빙과도 다름
}


@pytest.fixture
def sample_pipeline_inputs():
    journal = fxp.load_journal(str(SAMPLE_DIR / "분개장.xlsx"))
    schedule = pd.read_excel(SAMPLE_DIR / "명세서_외화자산부채명세서.xlsx")
    return journal, schedule


class TestSampleDataOracle:
    def test_answer_key_matches_expected_error_normal_split(self):
        """검증포인트_참고용.xlsx 자체가 이 테스트의 전제(어떤 거래가 오류/정상인지)와
        어긋나지 않는지 먼저 확인 - 샘플 데이터가 나중에 바뀌면 이 테스트가 먼저 깨진다."""
        answer_key = pd.read_excel(SAMPLE_DIR / "검증포인트_참고용.xlsx")
        answer_map = dict(zip(answer_key["거래ID"], answer_key["정상/오류"]))
        for txn in ["TXN002", "TXN003", "TXN004", "TXN007", "TXN008", "TXN013"]:
            assert answer_map[txn] == "오류", f"{txn}는 오류 케이스여야 함"
        for txn in ["TXN001", "TXN005", "TXN006", "TXN009", "TXN010", "TXN011", "TXN012", "TXN014"]:
            assert answer_map[txn] == "정상", f"{txn}는 정상 케이스여야 함"

    def test_settlement_screening_matches_expected_flags(self, sample_pipeline_inputs, monkeypatch):
        journal, _schedule = sample_pipeline_inputs
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch(SAMPLE_FAKE_RATES))

        settlements = fxp.extract_settlement_transactions(journal)
        screened = fxp.screen_fx_settlements(settlements).set_index("거래ID")

        # TXN001·TXN005: 실제 환율 그대로 정확히 계상 - 개별 스크리닝 통과(false positive 없음)
        assert screened.loc["TXN001", "1차플래그"] == "적정(스크리닝통과)"
        assert screened.loc["TXN005", "1차플래그"] == "적정(스크리닝통과)"
        # TXN002·TXN004: 문제가 있지만 괴리율 자체는 작아 1단계는 통과 (합계중요성에서 걸리는 설계)
        assert screened.loc["TXN002", "1차플래그"] == "적정(스크리닝통과)"
        assert screened.loc["TXN004", "1차플래그"] == "적정(스크리닝통과)"
        # TXN003: JPY 100단위 나누기 누락 - 괴리가 커서 1단계에서 바로 걸려야 함
        assert screened.loc["TXN003", "1차플래그"] == "이상치(정밀검증필요)"

        # TXN009: GBP 분할결제 2건 모두 실제 환율대로 정확히 반영 - 둘 다 통과해야 함
        assert (screened.loc["TXN009", "1차플래그"] == "적정(스크리닝통과)").all()
        assert len(screened.loc["TXN009"]) == 2
        # TXN011: CHF, 우대환율로 매매기준율과 괴리가 커서 1단계에서는 이상치로 걸려야 함
        # (2단계 증빙 OCR에서 최종적으로 적정으로 해소되는 건 TestOcrEvidenceForNewCurrencies에서 확인)
        assert screened.loc["TXN011", "1차플래그"] == "이상치(정밀검증필요)"
        # TXN012: CAD, 실제 환율 그대로 정확히 반영 - 통과해야 함
        assert screened.loc["TXN012", "1차플래그"] == "적정(스크리닝통과)"
        # TXN013: AUD, 회사가 실제 잘못된 환율을 적용해 1단계에서 이상치로 걸려야 함
        assert screened.loc["TXN013", "1차플래그"] == "이상치(정밀검증필요)"

        # 개별로는 안 걸려도, 합계로 보면 중요성을 초과해야 한다(README가 설명하는 설계 의도).
        agg = fxp.check_aggregate_materiality(screened.reset_index(), ampt=3000000)
        assert agg["중요성초과여부"]
        assert agg["확인불가건수"] == 0

    def test_reference_date_mismatch_matches_expected(self, sample_pipeline_inputs, monkeypatch):
        journal, _schedule = sample_pipeline_inputs
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch(SAMPLE_FAKE_RATES))

        settlements = fxp.extract_settlement_transactions(journal)
        ref_check = fxp.detect_reference_date_mismatch(settlements).set_index("거래ID")

        assert ref_check.loc["TXN001", "기준일판정"] == "정상(결제일 환율 사용 확인)"
        assert ref_check.loc["TXN002", "기준일판정"].startswith("기준일 오류 의심")
        assert ref_check.loc["TXN002", "추정사용일자"] == "2025-04-30"
        assert ref_check.loc["TXN005", "기준일판정"] == "정상(결제일 환율 사용 확인)"

    def test_yearend_translation_matches_expected(self, sample_pipeline_inputs, monkeypatch):
        journal, schedule = sample_pipeline_inputs
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch(SAMPLE_FAKE_RATES))

        ye_txns = fxp.extract_yearend_transactions(journal)
        missing_ye = fxp.get_unsettled_yearend_candidates(journal, schedule)
        ye_result = fxp.verify_yearend_translation(ye_txns, missing_ye, "2025-12-31").set_index("거래ID")

        assert ye_result.loc["TXN006", "일치여부"] == "적정"
        assert ye_result.loc["TXN008", "일치여부"] == "부적정(환율 오류)"
        assert ye_result.loc["TXN007", "일치여부"] == "부적정(기말평가 누락)"
        # TXN010(HKD)·TXN014(SGD) - 신규 통화도 기말 재평가가 정상적으로 통과해야 한다
        assert ye_result.loc["TXN010", "일치여부"] == "적정"
        assert ye_result.loc["TXN014", "일치여부"] == "적정"

    def test_ledger_reconciliation_catches_seeded_adjustment_gap(self, sample_pipeline_inputs):
        journal, _schedule = sample_pipeline_inputs
        ledger = pd.read_excel(SAMPLE_DIR / "계정별원장.xlsx")
        result = fxp.verify_ledger_reconciliation(journal, ledger).set_index("계정코드")

        # generate_samples.py가 907(외환차익) 계정원장 기말잔액에 50,000원 수기조정분을
        # 의도적으로 심어뒀으므로, 분개장 합계와 어긋나야 한다(완전성 이슈 탐지).
        assert result.loc["907", "일치여부"] == "부적정(분개장-원장 불일치)"
        assert result.loc["957", "일치여부"] == "적정"
        assert result.loc["908", "일치여부"] == "적정"
        assert result.loc["958", "일치여부"] == "적정"

    def test_ocr_evidence_resolves_txn011_and_txn013(self, sample_pipeline_inputs, monkeypatch):
        """TXN011(CHF, 우대환율)과 TXN013(AUD, 실제 오류)은 evidence/TXN011.png,
        evidence/TXN013.png가 실제로 존재하는 2단계 증빙 OCR 케이스다. anthropic
        클라이언트는 호출하지 않고 extract_rate_from_evidence만 목킹해서, 실제
        증빙 파일 탐색(_find_evidence_file)과 환율 일치/불일치 판정 로직만 확인한다
        - Claude Vision이 이미지를 실제로 정확히 읽는지는 tests/manual_ocr_check.py
        (수동, 실 API 호출)의 영역이다."""
        journal, _schedule = sample_pipeline_inputs
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch(SAMPLE_FAKE_RATES))
        monkeypatch.setattr(fxp, "EVIDENCE_DIR", str(ROOT_DIR / "evidence"))

        def _fake_extract(image_path):
            if "TXN011" in image_path:
                return {"거래일자": "2025-06-05", "통화": "CHF", "적용환율": 1750.00,
                         "외화금액": 6000.0, "원화금액": 10500000}
            if "TXN013" in image_path:
                return {"거래일자": "2025-08-14", "통화": "AUD", "적용환율": 918.75,
                         "외화금액": 25000.0, "원화금액": 22968750}
            raise AssertionError(f"예상치 못한 증빙 경로: {image_path}")

        monkeypatch.setattr(fxp, "extract_rate_from_evidence", _fake_extract)

        settlements = fxp.extract_settlement_transactions(journal)
        screened = fxp.screen_fx_settlements(settlements)
        flagged = screened[screened["1차플래그"] == "이상치(정밀검증필요)"]
        verified = fxp.verify_with_evidence(flagged).set_index("거래ID")

        # TXN011: 증빙 환율(1750.00)이 회사 계상 환율과 정확히 일치 -> 적정으로 해소
        assert verified.loc["TXN011", "증빙상태"] == "증빙 확인 완료"
        assert verified.loc["TXN011", "최종판정"] == "적정(증빙과 일치)"
        # TXN013: 증빙 환율(918.75)이 회사가 잘못 계상한 환율(980.00)과 불일치 -> 재확인
        assert verified.loc["TXN013", "증빙상태"] == "증빙 확인 완료"
        assert verified.loc["TXN013", "최종판정"] == "부적정(증빙과 불일치 - 재계산 필요)"


# ------------------------------------------------------------------
# 엑셀 내보내기 - 셀 단위 검증은 하지 않고, 예외 없이 예상 시트가 생성되는지만 확인
# ------------------------------------------------------------------

class TestExportResultsToExcel:
    def test_creates_workbook_with_expected_sheets(self, tmp_path):
        # 이번 기간에 결제 건이 하나도 없는 극단적인 경우(연중 미결제 건만 있는 회사 등)를
        # 실제 함수 체인으로 재현한다 - extract_settlement_transactions부터 시작해야
        # '결제일' 컬럼의 datetime dtype이 실제 파이프라인과 동일하게 유지된다.
        empty_journal = pd.DataFrame(columns=["전표번호(거래ID)", "일자", "구분", "계정코드",
                                               "계정과목", "차변금액", "대변금액", "통화",
                                               "외화금액", "적용환율", "거래처"])
        empty_settlements = fxp.extract_settlement_transactions(empty_journal)
        screened = fxp.screen_fx_settlements(empty_settlements)
        ref_check = fxp.detect_reference_date_mismatch(empty_settlements)
        verified = fxp.verify_with_evidence(screened)
        ye_result = fxp.verify_yearend_translation(
            pd.DataFrame(columns=["거래ID", "통화", "외화금액", "회사적용환율(기말)"]),
            pd.DataFrame(columns=["전표번호(거래ID)", "통화", "기말미결제외화잔액"]),
            "2025-12-31",
        )
        empty = pd.DataFrame()
        agg = {"AMPT": 3000000, "순차이합계(KRW)": 0, "절대값차이합계(KRW)": 0,
               "중요성초과여부": False, "판정": "전체 적정(합계 중요성 이내)", "확인불가건수": 0}
        output_path = tmp_path / "검증결과_테스트.xlsx"

        fxp.export_results_to_excel(
            str(output_path),
            counterparty_summary=empty, rate_trend=empty,
            screened=screened, agg=agg, ref_check=ref_check, verified=verified,
            ye_result=ye_result, ledger_result=empty, rollforward_result=empty,
            full_detail=empty,
        )

        assert output_path.exists()
        import openpyxl
        wb = openpyxl.load_workbook(output_path)
        assert set(wb.sheetnames) == {
            "안내", "요약", "A.표본전체상세", "B.분석적검토",
            "C.외환차손익_상세", "D.외화환산손익", "E.완전성검증",
        }


# ------------------------------------------------------------------
# 감사조서(Word) 자동생성 - 문서가 열리는지, 필수 섹션/Exception이 반영됐는지만 확인
# (셀 서식 등은 검증하지 않음, Excel과 마찬가지로 텍스트 존재 여부만 확인)
# ------------------------------------------------------------------

class TestGenerateAuditWorkpaper:
    def test_workpaper_reflects_sample_data_exceptions(self, tmp_path, sample_pipeline_inputs, monkeypatch):
        journal, schedule = sample_pipeline_inputs
        monkeypatch.setattr(fxp, "fetch_rates_for_date", _make_fake_fetch(SAMPLE_FAKE_RATES))
        # 증빙 폴더를 빈 임시경로로 돌려서, 2단계 OCR(실제 API 호출)이 아예 시도되지
        # 않도록 한다 - _find_evidence_file이 항상 None을 반환하므로 네트워크 호출 없음.
        monkeypatch.setattr(fxp, "EVIDENCE_DIR", str(tmp_path / "no_evidence"))
        ledger = pd.read_excel(SAMPLE_DIR / "계정별원장.xlsx")

        settlements = fxp.extract_settlement_transactions(journal)
        screened = fxp.screen_fx_settlements(settlements)
        ref_check = fxp.detect_reference_date_mismatch(settlements)
        agg = fxp.check_aggregate_materiality(screened, ampt=3000000)
        to_verify = (screened if agg["중요성초과여부"]
                     else screened[screened["1차플래그"] == "이상치(정밀검증필요)"])
        verified = fxp.verify_with_evidence(to_verify)

        ye_txns = fxp.extract_yearend_transactions(journal)
        missing_ye = fxp.get_unsettled_yearend_candidates(journal, schedule)
        ye_result = fxp.verify_yearend_translation(ye_txns, missing_ye, "2025-12-31")

        ledger_result = fxp.verify_ledger_reconciliation(journal, ledger)
        rollforward_result = fxp.build_rollforward_verification(journal, schedule, screened)

        output_path = tmp_path / "감사조서_테스트.docx"
        fxp.generate_audit_workpaper(
            str(output_path),
            screened=screened, agg=agg, ref_check=ref_check, verified=verified,
            ye_result=ye_result, ledger_result=ledger_result,
            rollforward_result=rollforward_result, year_end_date="2025-12-31",
        )

        assert output_path.exists()
        import docx
        doc = docx.Document(str(output_path))

        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        for expected in ["1. 검증 목적", "2. 검증 대상", "3. 수행 절차",
                          "4. 검증 결과", "5. 결론"]:
            assert expected in headings

        all_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                all_text += "\n" + " ".join(cell.text for cell in row.cells)

        # sample_data는 합계 중요성을 초과하도록 설계돼 있음(TestSampleDataOracle과 동일 전제)
        assert agg["중요성초과여부"]
        assert "중요한 차이가 발견되었다" in all_text
        # TXN003(JPY 100단위 나누기 누락)은 1단계에서 바로 걸리는 대표 이상치라
        # Exception 표에 반드시 나타나야 한다.
        assert "TXN003" in all_text
        # 계정별원장 대사에서 걸리는 외환차익(907) 계정 불일치도 나타나야 한다
        # (Exception 표의 식별자는 계정코드가 아닌 계정과목명을 사용).
        assert "외환차익" in all_text

    def test_conclusion_wording_when_no_material_difference(self, tmp_path):
        empty_journal = pd.DataFrame(columns=["전표번호(거래ID)", "일자", "구분", "계정코드",
                                               "계정과목", "차변금액", "대변금액", "통화",
                                               "외화금액", "적용환율", "거래처"])
        empty_settlements = fxp.extract_settlement_transactions(empty_journal)
        screened = fxp.screen_fx_settlements(empty_settlements)
        ref_check = fxp.detect_reference_date_mismatch(empty_settlements)
        verified = fxp.verify_with_evidence(screened)
        ye_result = fxp.verify_yearend_translation(
            pd.DataFrame(columns=["거래ID", "통화", "외화금액", "회사적용환율(기말)"]),
            pd.DataFrame(columns=["전표번호(거래ID)", "통화", "기말미결제외화잔액"]),
            "2025-12-31",
        )
        empty = pd.DataFrame()
        agg = {"AMPT": 3000000, "순차이합계(KRW)": 0, "절대값차이합계(KRW)": 0,
               "중요성초과여부": False, "판정": "전체 적정(합계 중요성 이내)", "확인불가건수": 0}

        output_path = tmp_path / "감사조서_정상.docx"
        fxp.generate_audit_workpaper(
            str(output_path),
            screened=screened, agg=agg, ref_check=ref_check, verified=verified,
            ye_result=ye_result, ledger_result=empty, rollforward_result=empty,
            year_end_date="2025-12-31",
        )

        import docx
        doc = docx.Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "중요한 차이는 발견되지 않았다" in all_text
        assert "Exception으로 식별된 건은 없다" in all_text
