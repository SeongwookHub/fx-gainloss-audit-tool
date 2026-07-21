# -*- coding: utf-8 -*-
"""
외환차손익 / 외화환산손익 검증 파이프라인 (1단계 스크리닝 + 2단계 증빙 OCR)
================================================================

[전체 구조]

  A. 외환차손익(결제 건, 분개장의 "결제" 라인)
     1단계: 전수 스크리닝 - 회사 적용환율 vs 그날 공식 매매기준율(수출입은행 API)
            괴리율이 5%를 넘는 건만 자동 추출 (전수조사가 불가능한 현실을 반영,
            "어떤 걸 표본으로 볼지"를 사람이 감으로 고르지 않고 수치 기준으로 자동 추출)
     2단계: 1단계에서 걸린 건만 증빙(은행 외환거래확인서/SWIFT/외화예금 거래명세표 등
            스캔 이미지)을 Claude Vision으로 읽어 실제 적용환율을 추출하고,
            그 환율로 재계산한 외환차손익을 회사 계상액과 최종 비교

  B. 외화환산손익(기말평가 건, 분개장의 "기말평가" 라인)
     결산일 하나의 공식 환율로 전수 재계산 (거래별로 증빙이 다를 이유가 없으므로
     OCR 2단계가 필요 없음 - 전수 자동 검증)

[사용 전 준비]
  1. data.go.kr에서 한국수출입은행 환율 API 인증키 발급 (EXIM_AUTH_KEY 환경변수)
  2. 증빙 이미지가 있다면 evidence/ 폴더에 "{거래ID}.png" 또는 "{거래ID}.jpg" 형태로 저장
     (예: evidence/TXN002.png) - 없으면 2단계는 "증빙 요청 필요"로 표시만 하고 넘어감
  3. Claude API 사용을 위해 ANTHROPIC_API_KEY 환경변수 설정 (2단계 OCR용)
  4. pip install -r requirements.txt (anthropic, pandas, openpyxl, requests, python-docx)

[실행]
  인자 없이 실행하면 sample_data/의 샘플 분개장·명세서·계정별원장을 대상으로 동작한다.
  --journal/--schedule/--ledger/--year-end-date/--output/--ampt로 실제 파일과 결산일을
  지정할 수 있다 (python fx_verification_pipeline.py --help 참고).
  --workpaper를 주면 검증결과.xlsx와 함께 감사조서 형식의 Word 문서(기본
  감사조서_외환차손익.docx, --workpaper-output으로 경로 변경 가능)도 생성한다.
"""

import os
import sys
import re
import json
import time
import math
import hashlib
import base64
import argparse
from datetime import datetime, timedelta

import pandas as pd
import requests
import openpyxl
from openpyxl.styles import Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from dotenv import load_dotenv

load_dotenv()  # 스크립트와 같은 폴더의 .env 파일에서 EXIM_AUTH_KEY/ANTHROPIC_API_KEY 등을 읽어옴

# Windows 콘솔 기본 코드페이지(cp949)는 "↳" 등 일부 문자를 인코딩하지 못해 print()가 죽으므로 stdout/stderr를 UTF-8로 고정
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        _stream.reconfigure(encoding="utf-8")

# ------------------------------------------------------------------
# 0. 설정
# ------------------------------------------------------------------

EXIM_AUTH_KEY = os.environ.get("EXIM_AUTH_KEY", "여기에_발급받은_인증키_입력")
EXIM_BASE_URL = "https://oapi.koreaexim.go.kr/site/program/financial/exchangeJSON"
# 참고: 2025.6.25부로 요청 URL 도메인이 www.koreaexim.go.kr -> oapi.koreaexim.go.kr로 변경됨.
# 기존 도메인(www.koreaexim.go.kr)은 점진적으로 종료 예정이라 응답이 없거나 타임아웃 날 수 있음.

TOLERANCE_PCT = 0.05          # 1단계 이상치 판단 기준: 괴리율 5%
GAIN_LOSS_TOLERANCE_KRW = 1000  # 재계산 금액과 회사 계상액의 허용 오차(원 단위 반올림 차이)
AMPT = float(os.environ.get("FX_AMPT", 3000000))  # 허용가능 오류금액(Tolerable Misstatement) - 감사팀 산정치로 교체

AR_ACCOUNT_CODE = "108"  # 외상매출금(채권)
AP_ACCOUNT_CODE = "251"  # 외상매입금(채무)

# 2단계 OCR이 "적정(증빙과 일치)"로 판정한 건 중 일부를 QC 표본으로 재확인하기 위한 비율.
# 거래 자체가 아니라 OCR 판정 로직의 신뢰성(위음성 여부)을 점검하는 목적.
OCR_RECHECK_SAMPLE_RATE = float(os.environ.get("FX_OCR_RECHECK_RATE", 0.15))
# 거래ID를 해시하는 솔트값. random 모듈의 전역/지역 시드가 아니라 거래ID 자체를 해시하므로,
# 매칭된 행의 순서나 개수가 달라져도(다른 기능 추가 등) 이미 뽑힌 거래ID의 표본 여부는 안 바뀐다.
OCR_RECHECK_SAMPLE_SEED = 42

CUR_UNIT_MAP = {
    "USD": "USD", "JPY": "JPY(100)", "EUR": "EUR", "CNH": "CNH", "CNY": "CNH",
    "GBP": "GBP", "HKD": "HKD", "CHF": "CHF", "CAD": "CAD", "AUD": "AUD", "SGD": "SGD",
}
# 수출입은행 API가 100단위로 고시하는 통화 (JPY 외 확인되는 대로 추가)
HUNDRED_UNIT_CURRENCIES = {"JPY"}

EVIDENCE_DIR = "evidence"      # 증빙 이미지 폴더
CLAUDE_MODEL = os.environ.get("FX_CLAUDE_MODEL", "claude-sonnet-4-6")  # 2단계 증빙 OCR용

_rate_cache = {}


# ------------------------------------------------------------------
# 1. 수출입은행 환율 API
# ------------------------------------------------------------------

def fetch_rates_for_date(date_str: str, _retries: int = 2) -> dict:
    """특정 날짜(YYYYMMDD)의 전체 통화 매매기준율표를 API에서 가져옴. RESULT 코드 체크 포함.
    일시적 타임아웃/연결 오류는 짧게 대기 후 재시도(최대 2회)."""
    if date_str in _rate_cache:
        return _rate_cache[date_str]

    params = {"authkey": EXIM_AUTH_KEY, "data": "AP01", "searchdate": date_str}

    for attempt in range(_retries + 1):
        try:
            resp = requests.get(EXIM_BASE_URL, params=params, timeout=10)
            resp.raise_for_status()
            data = resp.json()
            break
        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError):
            if attempt < _retries:
                time.sleep(1.5 * (attempt + 1))  # 1.5초, 3초 순으로 대기 후 재시도
                continue
            raise

    if not data:
        # 주말/공휴일 등으로 데이터가 없는 경우 (빈 리스트) - 폴백 대상
        _rate_cache[date_str] = {}
        return {}

    rate_table = {}
    for row in data:
        result = str(row.get("result", ""))
        if result == "2":
            raise ValueError("수출입은행 API 오류(RESULT=2): DATA 코드 오류")
        elif result == "3":
            raise ValueError("수출입은행 API 오류(RESULT=3): 인증코드(authkey) 오류 - 키를 확인하세요")
        elif result == "4":
            raise ValueError("수출입은행 API 오류(RESULT=4): 일일 호출 제한 초과 - 내일 다시 시도하거나 캐시를 활용하세요")
        try:
            rate_table[row["cur_unit"]] = float(row["deal_bas_r"].replace(",", ""))
        except (KeyError, ValueError):
            continue

    _rate_cache[date_str] = rate_table
    time.sleep(0.15)  # 짧은 pacing 지연 - 짧은 시간에 요청이 몰려 서버가 지연/거부하는 것을 예방
    return rate_table


def get_official_rate(date_obj: datetime, cur_unit: str, max_lookback_days: int = 5) -> tuple:
    """해당일 공식 매매기준율. 주말/공휴일이면 직전 영업일로 폴백.
    반환값: (환율, 실제 적용된 날짜) - 폴백된 경우 실제 날짜를 알아야 리포트에 남길 수 있음"""
    for i in range(max_lookback_days + 1):
        d = date_obj - timedelta(days=i)
        rate_table = fetch_rates_for_date(d.strftime("%Y%m%d"))
        if cur_unit in rate_table:
            return rate_table[cur_unit], d
    raise ValueError(f"{date_obj.date()} 기준 {cur_unit} 환율을 찾을 수 없습니다 (휴장일 {max_lookback_days}일 초과)")


def normalize_rate(raw_rate: float, currency: str) -> float:
    """100단위로 고시되는 통화(JPY 등)는 1단위당 단가로 환산."""
    return raw_rate / 100 if currency in HUNDRED_UNIT_CURRENCIES else raw_rate


def resolve_cur_unit(currency: str) -> tuple:
    """CUR_UNIT_MAP에 등록되지 않은 통화코드를 만나면, 원 코드를 그대로 시도는 하되
    (수출입은행 API 코드와 우연히 일치할 수도 있으므로) 경고 메시지를 함께 반환한다.
    호출부는 경고가 있으면 결과 행에 '통화코드 매핑 미등록'을 명확히 표시해야 한다."""
    if currency in CUR_UNIT_MAP:
        return CUR_UNIT_MAP[currency], None
    return currency, f"통화코드 매핑 미등록({currency}) - CUR_UNIT_MAP 확인 필요"


# ------------------------------------------------------------------
# 2. 분개장 파싱 - 결제 건 / 기말평가 건 추출
# ------------------------------------------------------------------

REQUIRED_JOURNAL_COLUMNS = ["전표번호(거래ID)", "일자", "구분", "계정코드", "계정과목"]


def _detect_header_row(path: str, max_scan_rows: int = 10) -> int:
    """실제 회사 파일에는 제목행/빈행이 표 위에 붙어있는 경우가 흔하다.
    필수 컬럼명이 실제로 나타나는 행을 찾아서 그 행을 헤더로 사용한다.
    못 찾으면 0(첫 행)을 그대로 반환해 기존 동작을 유지한다."""
    preview = pd.read_excel(path, header=None, nrows=max_scan_rows)
    for i in range(len(preview)):
        row_values = set(str(v).strip() for v in preview.iloc[i].tolist())
        if all(col in row_values for col in REQUIRED_JOURNAL_COLUMNS):
            return i
    return 0


def load_journal(path: str) -> pd.DataFrame:
    header_row = _detect_header_row(path)
    df = pd.read_excel(path, header=header_row)
    df = df.dropna(how="all")  # 완전히 빈 행 제거 (제목/구분용 빈 행 대응)

    # 계정코드가 "108.0", " 108" 등으로 지저분하게 들어와도 비교 가능하도록 정규화
    df["계정코드"] = df["계정코드"].apply(_normalize_account_code)
    # 계정과목명 앞뒤 공백 제거 (ERP export에서 흔함)
    df["계정과목"] = df["계정과목"].astype(str).str.strip()
    # 통화 코드 대소문자 통일
    df["통화"] = df["통화"].astype(str).str.strip().str.upper()
    # 외화금액이 "15,000" 같은 콤마 포함 텍스트로 들어온 경우 숫자로 변환
    if "외화금액" in df.columns:
        df["외화금액"] = (
            df["외화금액"].astype(str).str.replace(",", "", regex=False).replace("nan", None)
        )
        df["외화금액"] = pd.to_numeric(df["외화금액"], errors="coerce")

    df["일자"] = pd.to_datetime(df["일자"])
    return df


def _normalize_account_code(code) -> str:
    """108, "108", "108.0", " 108 " 등 다양한 형태로 들어오는 계정코드를
    일관된 문자열 형태("108")로 정규화."""
    try:
        return str(int(float(str(code).strip())))
    except (ValueError, TypeError):
        return str(code).strip()


SETTLEMENT_COLUMNS = [
    "거래ID", "결제일", "통화", "구분", "거래처", "데이터무결성경고",
    "외화금액", "결제원화금액", "발생원화금액", "회사적용환율(내재)", "회사계상_외환차익차손",
]


def extract_settlement_transactions(journal_df: pd.DataFrame) -> pd.DataFrame:
    """거래ID+결제일 단위로 '결제' 구분 라인을 묶어 결제 이벤트를 요약.
    한 거래가 여러 날짜에 나눠 결제된 경우(분할결제)에도 각 결제일을 별도
    이벤트로 인식한다 (거래ID만으로 묶으면 분할결제 중 일부가 누락되거나
    손익이 뒤섞이는 문제가 있었음)."""
    settle_rows = journal_df[(journal_df["구분"] == "결제") & (journal_df["외화금액"].notna())].copy()

    records = []
    for (txn_id, settle_date), group in settle_rows.groupby(["전표번호(거래ID)", "일자"]):
        # 채권/채무 결제 모두 '외화예금' 라인이 결제 시점 실제 자금 흐름의 기준
        cash_row = group[group["계정과목"] == "외화예금"]
        row = cash_row.iloc[0] if not cash_row.empty else group.iloc[0]
        currency = row["통화"]
        fc_amount = row["외화금액"]
        partner = row["거래처"]

        # 적용환율 '컬럼값'을 그대로 믿지 않고, 실제 분개된 원화금액에서 환율을 역산.
        booked_krw = row[["차변금액", "대변금액"]].fillna(0).abs().max()
        implied_rate = booked_krw / fc_amount if fc_amount else None

        # 907(외환차익)/957(외환차손) 계정 금액은 '같은 결제일 그룹' 안에서만 찾는다
        # (분할결제의 경우 결제 건마다 손익 라인이 따로 있으므로, 거래ID 전체를 보면
        #  서로 다른 결제 건의 손익이 뒤섞인다)
        # 907(외환차익)/957(외환차손) 라인은 외화금액이 비어있는(NaN) 경우가 많아
        # 앞서 외화금액 기준으로 필터링된 group(=settle_rows)에는 안 잡힐 수 있다.
        # journal_df 원본에서 같은 거래ID + 같은 결제일로 다시 조회해야 누락이 없다.
        same_leg = journal_df[(journal_df["전표번호(거래ID)"] == txn_id) &
                               (journal_df["구분"] == "결제") &
                               (journal_df["일자"] == settle_date)]
        gain_row = same_leg[same_leg["계정과목"] == "외환차익"]
        loss_row = same_leg[same_leg["계정과목"] == "외환차손"]

        booked_gain_loss = 0
        if not gain_row.empty:
            booked_gain_loss = gain_row.iloc[0][["차변금액", "대변금액"]].fillna(0).abs().max()
        elif not loss_row.empty:
            booked_gain_loss = -loss_row.iloc[0][["차변금액", "대변금액"]].fillna(0).abs().max()

        # 채권/채무 구분 - 거래ID 전체의 발생 라인 기준 (분할결제라도 발생은 1건)
        occur_rows = journal_df[(journal_df["전표번호(거래ID)"] == txn_id) &
                                 (journal_df["구분"] == "발생") &
                                 (journal_df["계정코드"].isin([AR_ACCOUNT_CODE, AP_ACCOUNT_CODE]))]
        side = "채권" if not occur_rows.empty and occur_rows.iloc[0]["계정코드"] == AR_ACCOUNT_CODE else "채무"

        # 데이터 무결성 경고: 같은 거래ID인데 통화나 거래처가 서로 다른 '발생' 라인이
        # 여러 개 섞여 있으면, 전표번호가 중복 채번되어 서로 다른 거래가 합쳐졌을 가능성이
        # 높다. 이 경우 아래 occur_krw는 첫 번째 라인 기준으로만 계산되므로 신뢰할 수 없다.
        integrity_warning = None
        if occur_rows["통화"].nunique() > 1 or occur_rows["거래처"].nunique() > 1:
            integrity_warning = (
                f"전표번호 중복 의심 - 동일 거래ID({txn_id})에 서로 다른 통화/거래처의 "
                f"발생 라인이 {len(occur_rows)}건 섞여 있음. 아래 금액은 신뢰할 수 없으니 원본 확인 필요"
            )

        # 발생원화금액(occur_krw)은 반드시 '원래 발생 전표'에서 가져와야 한다.
        # 결제 라인의 채권/채무 상계 금액을 그대로 쓰면, 그 상계 자체가 오류인 경우
        # (예: 외환차익을 인식하지 않고 결제금액으로 그대로 상계한 케이스) 오류가 있는
        # 금액을 '정상 발생액'으로 오인하게 되어 그 오류 자체를 놓치게 된다.
        # 분할결제는 원 발생액을 이번 결제 건의 외화금액 비중만큼 비례 안분한다.
        if not occur_rows.empty:
            total_occur_fc = occur_rows["외화금액"].sum()
            total_occur_krw = occur_rows.iloc[0][["차변금액", "대변금액"]].fillna(0).abs().max()
            if total_occur_fc:
                occur_krw = total_occur_krw * (fc_amount / total_occur_fc)
            else:
                occur_krw = total_occur_krw
        else:
            occur_krw = None

        records.append({
            "거래ID": txn_id, "결제일": settle_date, "통화": currency, "구분": side, "거래처": partner,
            "데이터무결성경고": integrity_warning,
            "외화금액": fc_amount, "결제원화금액": booked_krw, "발생원화금액": occur_krw,
            "회사적용환율(내재)": implied_rate,
            "회사계상_외환차익차손": booked_gain_loss,
        })

    if not records:
        # 이번 기간에 결제 건이 하나도 없어도(연중 미결제 건만 있는 회사 등), 다운스트림이
        # 기대하는 컬럼과 '결제일'의 datetime dtype을 유지해야 이후 merge에서 안전하다.
        empty = pd.DataFrame(columns=SETTLEMENT_COLUMNS)
        empty["결제일"] = pd.to_datetime(empty["결제일"])
        return empty

    return pd.DataFrame(records)


def extract_yearend_transactions(journal_df: pd.DataFrame) -> pd.DataFrame:
    """'기말평가' 구분 라인에서 거래ID별 재평가 내역 추출."""
    ye_rows = journal_df[(journal_df["구분"] == "기말평가") & (journal_df["외화금액"].notna())].copy()

    records = []
    for txn_id, group in ye_rows.groupby("전표번호(거래ID)"):
        row = group.iloc[0]
        records.append({
            "거래ID": txn_id, "결산일": row["일자"], "통화": row["통화"],
            "외화금액": row["외화금액"], "회사적용환율(기말)": row["적용환율"],
        })
    return pd.DataFrame(records)


def get_unsettled_yearend_candidates(journal_df: pd.DataFrame, schedule_df: pd.DataFrame) -> pd.DataFrame:
    """명세서(외화자산부채명세서) 기준 기말 미결제 건 중, 분개장에 기말평가 라인이
    아예 없는 거래를 찾아냄 (재평가 누락 탐지)."""
    outstanding = schedule_df[schedule_df["기말미결제외화잔액"] > 0]
    ye_done_ids = set(extract_yearend_transactions(journal_df)["거래ID"]) if not journal_df.empty else set()
    missing = outstanding[~outstanding["전표번호(거래ID)"].isin(ye_done_ids)]
    return missing


# ------------------------------------------------------------------
# 3. 1단계: 외환차손익 스크리닝
# ------------------------------------------------------------------

def screen_fx_settlements(settlements: pd.DataFrame) -> pd.DataFrame:
    """각 결제 건에 대해:
    1) 공식 매매기준율과의 '괴리율'을 계산해 5% 초과 건을 개별 플래그
    2) 공식환율 기준으로 재계산한 외환차손익과 회사 계상액의 '금액 차이(KRW)'도 함께 산출
       (개별로는 5% 이내라도, 이 금액 차이들을 합산해 중요성 검토를 하기 위함)
    회사측 환율은 실제 분개 금액에서 역산한 내재환율(원 단위, 정규화 불필요)을 사용.
    이번 기간에 결제 건이 하나도 없어도(연중 미결제 건만 있는 회사 등), 컬럼이 통째로
    빠진 DataFrame을 반환하면 다운스트림(엑셀 내보내기 등)이 깨지므로 컬럼 구조를 유지한다."""
    if settlements.empty:
        # concat으로 이어붙여야 settlements의 기존 컬럼 dtype(특히 '결제일'의 datetime64)이
        # 유지된다 - columns=[...]로 새로 만들면 전부 object dtype이 되어, 이후
        # build_fx_detail_report의 결제일 기준 merge에서 dtype 불일치 오류가 난다.
        extra = pd.DataFrame(columns=[
            "공식매매기준율", "공식환율기준일", "괴리율(%)", "1차플래그",
            "재계산_외환차손익(공식환율기준)", "회사계상액과의차이(KRW)", "비고",
        ])
        return pd.concat([settlements.reset_index(drop=True), extra], axis=1)

    results = []
    for _, row in settlements.iterrows():
        cur_unit, cur_warning = resolve_cur_unit(row["통화"])
        if cur_warning:
            # 통화코드가 CUR_UNIT_MAP에 없으면 API가 우연히 받아줄 수도 있는 값을 그대로
            # 조회하지 않고, detect_reference_date_mismatch/verify_yearend_translation과
            # 동일하게 조회 자체를 건너뛰고 확인불가로 남긴다.
            results.append({
                **row.to_dict(),
                "공식매매기준율": None, "공식환율기준일": None, "괴리율(%)": None,
                "1차플래그": f"확인불가({cur_warning})",
                "재계산_외환차손익(공식환율기준)": None, "회사계상액과의차이(KRW)": None,
                "비고": cur_warning,
            })
            continue
        try:
            official_raw, actual_date = get_official_rate(row["결제일"], cur_unit)
        except ValueError as e:
            # 환율을 못 찾은 거래 한 건 때문에 배치 전체가 죽지 않도록, 이 건만
            # '오류' 플래그로 남기고 나머지 거래는 계속 검증한다.
            results.append({
                **row.to_dict(),
                "공식매매기준율": None, "공식환율기준일": None, "괴리율(%)": None,
                "1차플래그": "오류(환율조회실패)",
                "재계산_외환차손익(공식환율기준)": None, "회사계상액과의차이(KRW)": None,
                "비고": cur_warning or str(e),
            })
            continue

        official_rate = normalize_rate(official_raw, row["통화"])
        company_rate = row["회사적용환율(내재)"]

        deviation_pct = abs(company_rate - official_rate) / official_rate
        flagged = deviation_pct > TOLERANCE_PCT

        # 공식환율로 재계산한 결제원화금액 및 외환차손익 (부호는 booked_gain_loss와 동일 규약:
        # 채권은 (결제-발생)이 이익, 채무는 (발생-결제)가 이익)
        recalculated_settle_krw = row["외화금액"] * official_rate
        if row["구분"] == "채권":
            recalculated_gain_loss = recalculated_settle_krw - row["발생원화금액"]
        else:
            recalculated_gain_loss = row["발생원화금액"] - recalculated_settle_krw

        diff_krw = recalculated_gain_loss - row["회사계상_외환차익차손"]

        results.append({
            **row.to_dict(),
            "공식매매기준율": round(official_rate, 4),
            "공식환율기준일": actual_date.strftime("%Y-%m-%d"),
            "괴리율(%)": round(deviation_pct * 100, 2),
            "1차플래그": "이상치(정밀검증필요)" if flagged else "적정(스크리닝통과)",
            "재계산_외환차손익(공식환율기준)": round(recalculated_gain_loss),
            "회사계상액과의차이(KRW)": round(diff_krw),
            "비고": cur_warning,
        })
    return pd.DataFrame(results)


# ------------------------------------------------------------------
# 3-보조0. 거래처별 연간 요약 & 환율 추이 (분석적 절차 - 드릴다운 전 1차 검토용)
# ------------------------------------------------------------------
#
# 실제 감사에서는 분개장을 한 줄씩 보기 전에, 거래처별로 연간 외환차손익을
# 집계한 요약본과 연중 환율 추이를 먼저 보고 "이상해 보이는" 거래처/시점만
# 골라 분개장을 드릴다운한다. 아래 두 함수가 그 1차 검토 단계를 지원한다.

def build_counterparty_summary(screened: pd.DataFrame) -> pd.DataFrame:
    """거래처(+통화) 단위로 연간 결제 건수, 총 외환차손익, 회사계상액과의 차이,
    평균/최대 괴리율, 이상치 건수를 집계. 요약본에서 숫자가 튀는 거래처만
    골라 분개장을 드릴다운하면 된다."""
    if screened.empty:
        return pd.DataFrame()

    def _agg(g):
        return pd.Series({
            "결제건수": len(g),
            "총외화금액": g["외화금액"].sum(),
            "회사계상_외환차손익_합계": g["회사계상_외환차익차손"].sum(),
            "재계산_외환차손익_합계": g["재계산_외환차손익(공식환율기준)"].sum(),
            "차이_합계(KRW)": g["회사계상액과의차이(KRW)"].sum(),
            "평균괴리율(%)": round(g["괴리율(%)"].mean(), 2),
            "최대괴리율(%)": round(g["괴리율(%)"].max(), 2),
            "이상치건수": (g["1차플래그"] == "이상치(정밀검증필요)").sum(),
        })

    summary = screened.groupby(["거래처", "통화"], dropna=False).apply(_agg, include_groups=False).reset_index()
    return summary.sort_values("차이_합계(KRW)", key=abs, ascending=False)


def build_rate_trend_summary(screened: pd.DataFrame) -> pd.DataFrame:
    """월별로 회사 평균 적용환율(내재)과 공식 평균환율을 나란히 보여줘, 연중
    환율 흐름에서 이상한 시점(월)이 있는지 한눈에 볼 수 있게 한다."""
    if screened.empty:
        return pd.DataFrame()

    df = screened.copy()
    df["결제월"] = pd.to_datetime(df["결제일"]).dt.strftime("%Y-%m")

    def _agg(g):
        return pd.Series({
            "거래건수": len(g),
            "회사평균적용환율": round(g["회사적용환율(내재)"].mean(), 2),
            "공식평균환율": round(g["공식매매기준율"].mean(), 4),
            "평균괴리율(%)": round(g["괴리율(%)"].mean(), 2),
            "이상치건수": (g["1차플래그"] == "이상치(정밀검증필요)").sum(),
        })

    trend = df.groupby(["결제월", "통화"], dropna=False).apply(_agg, include_groups=False).reset_index()
    return trend.sort_values(["통화", "결제월"])


def check_aggregate_materiality(screened: pd.DataFrame, ampt: float) -> dict:
    """개별 건이 전부 5% 이내로 통과했더라도, 재계산액과 회사계상액의 차이를
    전체 합산했을 때 AMPT(허용가능 오류금액)를 넘는지 확인.
    순합계(넷)와 절대값합계(그로스) 둘 다 보여줌 - 방향이 서로 반대인 오류가
    상쇄되어 순합계는 작아 보여도 그로스 기준으로는 클 수 있기 때문.
    환율조회 실패 등으로 차이금액을 못 구한 행(None)은 합계에서 제외하고
    건수만 별도로 집계한다 - 조용히 누락시키지 않고 존재를 드러내기 위함."""
    diffs = pd.to_numeric(screened["회사계상액과의차이(KRW)"], errors="coerce")
    unresolved_count = int(diffs.isna().sum())
    net_sum = diffs.sum()
    gross_sum = diffs.abs().sum()
    breach = abs(net_sum) > ampt or gross_sum > ampt

    return {
        "AMPT": ampt,
        "순차이합계(KRW)": round(net_sum),
        "절대값차이합계(KRW)": round(gross_sum),
        "중요성초과여부": breach,
        "판정": "전체 재검토 필요(합계 중요성 초과)" if breach else "전체 적정(합계 중요성 이내)",
        "확인불가건수": unresolved_count,
    }


# ------------------------------------------------------------------
# 3-보조. 결제일-환율기준일 불일치 탐지
# ------------------------------------------------------------------
#
# 왜 필요한가: 5% 괴리율 기준(크기 기준)은 "전월말 환율을 잘못 쓴" 것 같은 오류를
# 놓칠 수 있다 (원/달러가 한 달 새 5% 넘게 안 움직이면 개별 판정을 통과해버림 -
# TXN002 사례에서 실측). 그래서 크기가 아니라 "회사가 쓴 환율이 실제로는 다른
# 날짜의 공식 환율과 정확히 일치하는가"를 별도로 확인한다. 소액이라도 날짜
# 자체를 잘못 쓴 경우라면 이 체크에서 걸린다.

REF_DATE_MATCH_TOLERANCE = 0.0005  # 환율이 '일치'한다고 볼 허용오차(0.05%) - API 반올림 오차 흡수용


def _candidate_wrong_dates(settle_date: datetime) -> list:
    """실무에서 실제로 자주 발생하는 '기준일 착오' 패턴 후보만 골라서 반환.
    (임의로 40일을 다 훑는 대신, 흔한 실수 패턴만 확인 -> API 호출을 대폭 절감)"""
    candidates = []

    # 전월말(직전월 마지막 날)
    first_of_this_month = settle_date.replace(day=1)
    prev_month_end = first_of_this_month - timedelta(days=1)
    candidates.append(prev_month_end)

    # 전영업일들 (T-1 ~ T-3, 주말 포함해서 최대 5일 전까지)
    for i in range(1, 6):
        candidates.append(settle_date - timedelta(days=i))

    # 1주일 전, 2주일 전 (담당자가 "지난주 환율로" 착각하는 경우)
    candidates.append(settle_date - timedelta(days=7))
    candidates.append(settle_date - timedelta(days=14))

    # 중복 제거, 결제일 이후 날짜는 제외
    seen = set()
    result = []
    for d in candidates:
        key = d.strftime("%Y%m%d")
        if key not in seen and d < settle_date:
            seen.add(key)
            result.append(d)
    return result


def detect_reference_date_mismatch(settlements: pd.DataFrame) -> pd.DataFrame:
    """회사가 사용한 환율이 결제일이 아닌 다른 '흔히 착각하는' 날짜의 공식환율과
    더 정확히 일치하는지 확인. 일치하는 과거 날짜를 찾으면 '기준일 오류 의심'으로
    표시하고 그 날짜를 함께 보여준다. 5% 이내로 통과한 건도 전부 검사 대상.
    (거래당 API 호출을 최대 8회 내외로 제한 - 이전 버전은 최대 40회까지 순차 호출해서
    수출입은행 서버의 사실상 레이트리밋에 걸렸었음)"""
    if settlements.empty:
        return pd.DataFrame(columns=["거래ID", "결제일", "회사적용환율(내재)", "기준일판정", "추정사용일자"])

    results = []
    for _, row in settlements.iterrows():
        cur_unit, cur_warning = resolve_cur_unit(row["통화"])
        company_rate = row["회사적용환율(내재)"]
        settle_date = row["결제일"]

        if cur_warning:
            results.append({
                "거래ID": row["거래ID"], "결제일": settle_date.strftime("%Y-%m-%d"),
                "회사적용환율(내재)": company_rate, "기준일판정": f"확인불가({cur_warning})",
                "추정사용일자": None,
            })
            continue

        matched_date = None
        for candidate_date in _candidate_wrong_dates(settle_date):
            try:
                rate_table = fetch_rates_for_date(candidate_date.strftime("%Y%m%d"))
            except ValueError:
                # 이 후보 날짜만 건너뛴다 - 바로 다음의 get_official_rate 호출에서
                # 같은 원인(키 오류 등)이면 그때 이 거래 전체를 '확인불가'로 남긴다.
                continue
            if cur_unit not in rate_table:
                continue
            candidate_rate = normalize_rate(rate_table[cur_unit], row["통화"])
            if candidate_rate == 0:
                continue
            if abs(company_rate - candidate_rate) / candidate_rate <= REF_DATE_MATCH_TOLERANCE:
                matched_date = candidate_date
                break

        try:
            official_settle_raw, _ = get_official_rate(settle_date, cur_unit)
        except ValueError as e:
            results.append({
                "거래ID": row["거래ID"], "결제일": settle_date.strftime("%Y-%m-%d"),
                "회사적용환율(내재)": company_rate, "기준일판정": f"확인불가(환율조회실패: {e})",
                "추정사용일자": None,
            })
            continue
        official_settle_rate = normalize_rate(official_settle_raw, row["통화"])
        already_matches_settle_date = (
            official_settle_rate != 0 and
            abs(company_rate - official_settle_rate) / official_settle_rate <= REF_DATE_MATCH_TOLERANCE
        )

        if already_matches_settle_date:
            verdict = "정상(결제일 환율 사용 확인)"
            matched_date_str = None
        elif matched_date is not None:
            verdict = f"기준일 오류 의심 - {matched_date.strftime('%Y-%m-%d')} 환율을 사용한 것으로 보임"
            matched_date_str = matched_date.strftime("%Y-%m-%d")
        else:
            verdict = "불일치(흔한 패턴과 매칭 안 됨 - 은행 우대환율 등 개별 사유 가능, 증빙 확인 필요)"
            matched_date_str = None

        results.append({
            "거래ID": row["거래ID"], "결제일": settle_date.strftime("%Y-%m-%d"),
            "회사적용환율(내재)": company_rate, "기준일판정": verdict, "추정사용일자": matched_date_str,
        })

    return pd.DataFrame(results)


# ------------------------------------------------------------------
# 4. 2단계: 증빙 OCR (Claude Vision)
# ------------------------------------------------------------------

def _find_evidence_file(txn_id: str) -> str | None:
    for ext in (".png", ".jpg", ".jpeg", ".pdf"):
        candidate = os.path.join(EVIDENCE_DIR, f"{txn_id}{ext}")
        if os.path.exists(candidate):
            return candidate
    return None


def extract_rate_from_evidence(image_path: str) -> dict:
    """증빙 이미지(은행 외환거래확인서 등)에서 실제 적용환율/금액을 Claude Vision으로 추출.
    반환: {"거래일자":..., "통화":..., "적용환율":..., "외화금액":..., "원화금액":...}
    파싱 실패 시 raw_text에 원본 응답을 담아 반환하니, 로그로 확인해서 프롬프트를 조정하세요."""
    import anthropic  # 로컬 실행 시 pip install anthropic 필요

    client = anthropic.Anthropic()  # ANTHROPIC_API_KEY 환경변수 사용

    with open(image_path, "rb") as f:
        image_b64 = base64.standard_b64encode(f.read()).decode("utf-8")

    ext = os.path.splitext(image_path)[1].lower()
    media_type = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg"}.get(ext.strip("."), "image/png")

    if ext == ".pdf":
        content_block = {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": image_b64}}
    else:
        content_block = {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": image_b64}}

    prompt = (
        "이 이미지는 은행의 외환거래확인서, SWIFT 통지서, 또는 외화예금 거래명세표입니다. "
        "여기서 실제 적용된 환율과 거래 정보를 추출하세요. "
        "다른 설명 없이 아래 형식의 JSON 객체 하나만 답하세요:\n"
        '{"거래일자": "YYYY-MM-DD", "통화": "USD", "적용환율": 1234.5, '
        '"외화금액": 10000, "원화금액": 12345000}\n'
        "값을 찾을 수 없는 항목은 null로 표기하세요."
    )

    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=500,
        messages=[{"role": "user", "content": [content_block, {"type": "text", "text": prompt}]}],
    )

    raw_text = "".join(b.text for b in message.content if b.type == "text")
    cleaned = re.sub(r"```json|```", "", raw_text).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return {"error": "JSON 파싱 실패", "raw_text": raw_text}


def _ocr_recheck_hash_key(transaction_id, seed: int) -> str:
    """거래ID를 sha256으로 해시해 정렬 키를 만든다. 내장 hash()는 PYTHONHASHSEED에 따라
    실행마다 값이 달라져 재현성이 깨지므로 쓰지 않는다."""
    return hashlib.sha256(f"{seed}:{transaction_id}".encode("utf-8")).hexdigest()


def _select_ocr_recheck_sample(rows: list) -> None:
    """OCR이 '적정(증빙과 일치)'로 판정한 행 중 일부를 골라 'OCR재확인표본' 플래그를 붙인다.
    '불일치' 판정 건은 이미 Exception으로 전수 확인 대상이 되므로 손대지 않는다 - 여기서
    보는 것은 거래 자체의 적정성이 아니라, OCR이 실제로는 불일치인데 일치로 오판(위음성)하지
    않았는지를 감사인이 표본으로 재확인하기 위한 QC 절차다. 거래ID를 해시해 순위를 매기고
    상위 N개를 뽑는다(거래ID 기반 고정) - random.Random(seed).sample()처럼 매칭된 행의
    리스트 순서/구성에 기대는 방식이면, 다른 기능이 추가되어 순서나 건수가 달라질 때 이미
    뽑혔던 거래ID의 표본 여부까지 흔들릴 수 있어 피한다. 문구는 FLAG_KEYWORDS와 겹치지
    않게 골라 Exception 하이라이트에 섞이지 않도록 한다."""
    matched = [r for r in rows if r.get("최종판정") == "적정(증빙과 일치)"]
    sampled_ids = set()
    if matched:
        sample_size = min(len(matched), max(1, math.ceil(len(matched) * OCR_RECHECK_SAMPLE_RATE)))
        ranked = sorted(matched, key=lambda r: _ocr_recheck_hash_key(r.get("거래ID"), OCR_RECHECK_SAMPLE_SEED))
        sampled_ids = {id(r) for r in ranked[:sample_size]}
    # 매칭된 행이 0건이어도 컬럼 자체는 항상 채워야, 호출부가 dtype/컬럼 접근 시 KeyError가 안 난다.
    for r in rows:
        r["OCR재확인표본"] = "재확인 대상(QC 표본)" if id(r) in sampled_ids else None


def verify_with_evidence(flagged_df: pd.DataFrame) -> pd.DataFrame:
    """1단계에서 플래그된 건에 대해 증빙 파일이 있으면 OCR로 검증, 없으면 '증빙요청필요' 표시.
    정밀검증 대상이 0건(이상치가 하나도 없는 정상적인 기간)이어도, 호출부가 기대하는
    컬럼이 빠지지 않도록 빈 DataFrame에도 컬럼 구조를 유지해서 반환한다."""
    if flagged_df.empty:
        return pd.DataFrame(columns=list(flagged_df.columns) + ["증빙상태", "증빙확인환율", "최종판정", "OCR재확인표본"])

    rows = []
    for _, row in flagged_df.iterrows():
        record = row.to_dict()
        evidence_path = _find_evidence_file(row["거래ID"])

        if evidence_path is None:
            record["증빙상태"] = "증빙 요청 필요"
            record["증빙확인환율"] = None
            record["최종판정"] = "미확인(증빙 미확보)"
            rows.append(record)
            continue

        try:
            extracted = extract_rate_from_evidence(evidence_path)
        except Exception as e:
            # ANTHROPIC_API_KEY 미설정/만료, 네트워크 오류 등으로 OCR 호출 자체가 실패해도
            # 이 거래만 '수기 확인 필요'로 남기고 나머지 거래 검증은 계속 진행한다.
            record["증빙상태"] = "OCR 호출 실패 - 수기 확인 필요"
            record["증빙확인환율"] = None
            record["최종판정"] = f"미확인(OCR 호출 오류: {e})"
            rows.append(record)
            continue

        if "error" in extracted:
            record["증빙상태"] = "OCR 인식 실패 - 수기 확인 필요"
            record["증빙확인환율"] = None
            record["최종판정"] = "미확인(OCR 실패)"
            rows.append(record)
            continue

        confirmed_rate = extracted.get("적용환율")
        record["증빙상태"] = "증빙 확인 완료"
        record["증빙확인환율"] = confirmed_rate

        if confirmed_rate is not None:
            currency = row["통화"]
            fc = row["외화금액"]
            # 재계산: 발생시점 원가와의 차이는 별도 원장 대조가 필요하므로,
            # 여기서는 "증빙 환율 vs 회사 계상 환율" 자체의 일치 여부를 우선 확인
            company_rate = row["회사적용환율(내재)"]
            rate_match = abs(confirmed_rate - company_rate) <= max(company_rate * 0.001, 0.01)
            record["최종판정"] = "적정(증빙과 일치)" if rate_match else "부적정(증빙과 불일치 - 재계산 필요)"
        else:
            record["최종판정"] = "미확인(증빙에서 환율 추출 실패)"

        rows.append(record)

    _select_ocr_recheck_sample(rows)
    return pd.DataFrame(rows)


# ------------------------------------------------------------------
# 5. 외화환산손익 (기말평가) - 전수 자동 검증
# ------------------------------------------------------------------

YEAREND_RESULT_COLUMNS = ["거래ID", "결산일", "통화", "외화금액", "회사적용환율(기말)", "공식결산환율", "일치여부"]


def verify_yearend_translation(ye_df: pd.DataFrame, missing_df: pd.DataFrame, year_end_date: str) -> pd.DataFrame:
    if ye_df.empty and missing_df.empty:
        return pd.DataFrame(columns=YEAREND_RESULT_COLUMNS)

    rows = []
    ye_dt = pd.to_datetime(year_end_date)

    for _, row in ye_df.iterrows():
        cur_unit, cur_warning = resolve_cur_unit(row["통화"])
        if cur_warning:
            rows.append({**row.to_dict(), "공식결산환율": None, "일치여부": f"확인불가({cur_warning})"})
            continue
        try:
            official_raw, actual_date = get_official_rate(ye_dt, cur_unit)
        except ValueError as e:
            rows.append({**row.to_dict(), "공식결산환율": None, "일치여부": f"확인불가(환율조회실패: {e})"})
            continue
        official_rate = normalize_rate(official_raw, row["통화"])
        company_rate = normalize_rate(row["회사적용환율(기말)"], row["통화"])

        rate_match = abs(company_rate - official_rate) <= official_rate * 0.001
        rows.append({
            **row.to_dict(),
            "공식결산환율": round(official_rate, 4),
            "일치여부": "적정" if rate_match else "부적정(환율 오류)",
        })

    for _, row in missing_df.iterrows():
        rows.append({
            "거래ID": row["전표번호(거래ID)"], "결산일": year_end_date, "통화": row["통화"],
            "외화금액": row["기말미결제외화잔액"], "회사적용환율(기말)": None,
            "공식결산환율": None, "일치여부": "부적정(기말평가 누락)",
        })

    return pd.DataFrame(rows)


# ------------------------------------------------------------------
# 5-보조. 계정별원장 대사 (분개장 <-> 계정별원장 교차검증)
# ------------------------------------------------------------------
#
# 분개장과 명세서만 봐서는 못 잡는 유형의 오류가 있다: 상위 시스템에서 수기로
# 조정분개를 넣었는데 그게 상세 분개장에는 반영이 안 된 경우. 분개장 자체는
# 내부적으로 앞뒤가 맞으니 분개장만 보면 이상이 없어 보이지만, 계정별원장(총계정원장)의
# 기말잔액과 대조하면 차이가 드러난다. 그래서 외환차손익/외화환산손익 관련
# 4개 계정(외환차익/외환차손/외화환산이익/외화환산손실)에 한해 분개장에서 집계한
# 차변·대변 합계를 계정별원장 상 금액과 대사한다.

LEDGER_RECONCILE_ACCOUNTS = {
    "907": "외환차익", "957": "외환차손", "908": "외화환산이익", "958": "외화환산손실",
}
LEDGER_RECONCILE_TOLERANCE_KRW = 1000  # 반올림 오차 흡수용


def verify_ledger_reconciliation(journal_df: pd.DataFrame, ledger_df: pd.DataFrame) -> pd.DataFrame:
    """분개장에서 계정별로 집계한 차변/대변 합계가 계정별원장(총계정원장) 상
    차변합계/대변합계와 일치하는지 확인. 코드/컬럼명이 문자열-숫자 혼재 등으로
    지저분해도 최대한 강건하게 매칭하도록 계정코드를 문자열로 정규화해서 비교."""
    ledger_df = ledger_df.copy()
    ledger_df["계정코드"] = ledger_df["계정코드"].astype(str).str.strip()

    rows = []
    for code, name in LEDGER_RECONCILE_ACCOUNTS.items():
        journal_rows = journal_df[journal_df["계정코드"].astype(str).str.strip() == code]
        journal_debit = pd.to_numeric(journal_rows["차변금액"], errors="coerce").fillna(0).sum()
        journal_credit = pd.to_numeric(journal_rows["대변금액"], errors="coerce").fillna(0).sum()

        ledger_row = ledger_df[ledger_df["계정코드"] == code]
        if ledger_row.empty:
            rows.append({
                "계정코드": code, "계정과목": name,
                "분개장_차변합계": round(journal_debit), "분개장_대변합계": round(journal_credit),
                "원장_차변합계": None, "원장_대변합계": None,
                "일치여부": "부적정(계정별원장에 해당 계정 없음)",
            })
            continue

        ledger_debit = pd.to_numeric(ledger_row.iloc[0]["차변합계"], errors="coerce")
        ledger_credit = pd.to_numeric(ledger_row.iloc[0]["대변합계"], errors="coerce")

        debit_match = abs(journal_debit - ledger_debit) <= LEDGER_RECONCILE_TOLERANCE_KRW
        credit_match = abs(journal_credit - ledger_credit) <= LEDGER_RECONCILE_TOLERANCE_KRW

        rows.append({
            "계정코드": code, "계정과목": name,
            "분개장_차변합계": round(journal_debit), "분개장_대변합계": round(journal_credit),
            "원장_차변합계": round(ledger_debit), "원장_대변합계": round(ledger_credit),
            "일치여부": "적정" if (debit_match and credit_match) else "부적정(분개장-원장 불일치)",
        })

    return pd.DataFrame(rows)


# ------------------------------------------------------------------
# 5-보조2. 계정매핑 (회사·현장·국가·본지점마다 다른 계정명 대응)
# ------------------------------------------------------------------
#
# 회사(특히 건설업처럼 현장/국가/본지점별로 계정을 쪼개는 업종)마다 실제
# 계정명이 다 다르다. "외환차익"이라는 표준 문자열이 그대로 안 나오는 경우가
# 흔하므로, 감사인이 회사의 계정과목 목록(시산표 등)을 한 번 넘겨주면
# 키워드로 1차 자동분류하고, 애매한 것만 확인받는 방식으로 처리한다.
# 확인이 끝난 매핑표는 이후 전 회계기간에 재사용 가능하다.

FX_ACCOUNT_KEYWORDS = {
    AR_ACCOUNT_CODE: (["외상매출금", "매출채권", "수출채권", "AR"], ["외화", "외환", "FX", "USD", "EUR", "JPY"]),
    AP_ACCOUNT_CODE: (["외상매입금", "매입채무", "수입채무", "AP"], ["외화", "외환", "FX", "USD", "EUR", "JPY"]),
    "103": (["외화예금", "외환예금"], []),
    "907": (["외환차익", "환차익", "외화차익", "FX GAIN", "FX 차익"], []),
    "957": (["외환차손", "환차손", "외화차손", "FX LOSS", "FX 차손"], []),
    "908": (["외화환산이익", "환산이익", "외화평가이익", "FX TRANSLATION GAIN"], []),
    "958": (["외화환산손실", "환산손실", "외화평가손실", "FX TRANSLATION LOSS"], []),
}
STANDARD_ACCOUNT_NAME = {
    AR_ACCOUNT_CODE: "외상매출금(외화)", AP_ACCOUNT_CODE: "외상매입금(외화)", "103": "외화예금",
    "907": "외환차익", "957": "외환차손", "908": "외화환산이익", "958": "외화환산손실",
}


def build_account_mapping(company_accounts: pd.DataFrame) -> pd.DataFrame:
    """company_accounts: ['회사계정코드', '회사계정명'] 컬럼을 가진 계정과목 목록
    (시산표, 계정과목 리스트 등)을 받아, 키워드 기준으로 표준분류를 1차 추천한다.
    하나의 표준분류에만 걸리면 자동 확정, 0개 또는 2개 이상 걸리면 확인 필요로 표시."""
    rows = []
    for _, row in company_accounts.iterrows():
        name_upper = str(row["회사계정명"]).upper()
        matched = []
        for std_code, (must_have, nice_to_have) in FX_ACCOUNT_KEYWORDS.items():
            if any(kw.upper() in name_upper for kw in must_have):
                matched.append(std_code)

        if len(matched) == 1:
            status = "자동확정"
            std_code = matched[0]
        elif len(matched) == 0:
            status = "확인필요(매칭 없음)"
            std_code = None
        else:
            status = f"확인필요(중복매칭 {matched})"
            std_code = None

        rows.append({
            "회사계정코드": row["회사계정코드"], "회사계정명": row["회사계정명"],
            "추천표준분류": std_code, "표준분류명": STANDARD_ACCOUNT_NAME.get(std_code),
            "상태": status,
        })
    return pd.DataFrame(rows)


def apply_account_mapping(journal_df: pd.DataFrame, confirmed_mapping: dict) -> pd.DataFrame:
    """감사인이 확인을 마친 매핑({회사계정코드: 표준코드, ...})을 분개장에 적용해
    계정코드/계정과목을 표준 값(108/251/103/907/957/908/958, 표준 계정명)으로
    치환한다. 이후 파이프라인 나머지 코드는 전혀 수정 없이 그대로 동작한다."""
    df = journal_df.copy()
    df["원본계정코드"] = df["계정코드"]
    df["원본계정과목"] = df["계정과목"]

    def _map_code(code):
        return confirmed_mapping.get(str(code).strip(), code)

    df["계정코드"] = df["계정코드"].apply(_map_code)
    df["계정과목"] = df["계정코드"].map(STANDARD_ACCOUNT_NAME).fillna(df["계정과목"])
    return df


# ------------------------------------------------------------------
# 5-보조3. 전표번호 자동 부여 (차변=대변이 되는 분개 덩어리 단위)
# ------------------------------------------------------------------
#
# 전표번호가 아예 없는 원본 분개장도 있다. 이 경우 원본 순서대로 훑으면서
# 누적(차변-대변)이 0으로 돌아오는 지점까지를 전표 1건으로 묶는다.
# 이건 순수 회계 기계 규칙(차변합계=대변합계)이라 사람 판단이 필요 없다.

def auto_assign_voucher_number(df: pd.DataFrame, tolerance: float = 10) -> pd.DataFrame:
    """'전표번호(거래ID)' 컬럼이 없는 분개장에 차변=대변=0 단위로 전표번호를
    자동 부여한다. 원본 파일에 있는 행 순서를 그대로 신뢰한다(전표의 차/대변
    라인이 서로 떨어져 뒤섞여 있는 경우는 지원하지 않음)."""
    df = df.reset_index(drop=True).copy()
    voucher_ids = []
    voucher_no = 1
    running_balance = 0.0
    lines_in_current = 0

    for _, row in df.iterrows():
        debit = row.get("차변금액")
        credit = row.get("대변금액")
        debit = 0 if pd.isna(debit) else debit
        credit = 0 if pd.isna(credit) else credit
        running_balance += (debit - credit)
        lines_in_current += 1
        voucher_ids.append(f"AUTO{voucher_no:05d}")

        if lines_in_current >= 2 and abs(running_balance) <= tolerance:
            voucher_no += 1
            running_balance = 0.0
            lines_in_current = 0

    if lines_in_current != 0:
        raise ValueError(
            f"마지막 전표(AUTO{voucher_no:05d})가 차변=대변으로 안 맞아떨어집니다 "
            f"(잔액 {running_balance}). 원본 분개장 순서나 누락된 라인을 확인하세요."
        )

    df["전표번호(거래ID)"] = voucher_ids
    return df


# ------------------------------------------------------------------
# 5-보조4. 거래처별 롤포워드(잔액 증감) 검증 - 거래ID 매칭 없이도 가능한 총액 검증
# ------------------------------------------------------------------
#
# "100만 달러를 10번에 나눠 갚은" 것처럼 발생-결제를 1:1로 짝짓기 어려운
# 경우에도, 거래처 단위 잔액 증감(기초+발생-기말=소멸장부가)만으로 실현
# 외환차손익 총액을 구할 수 있다 (재고자산 롤포워드와 같은 원리).
# 이 총액을 거래ID 매칭 기반 합계(bottom-up)와 비교하면, 매칭 자체가 안 맞는
# 경우(거래 누락/중복)를 거래ID 매칭 없이도 발견할 수 있다.

def build_rollforward_verification(journal_df: pd.DataFrame, schedule_df: pd.DataFrame,
                                    screened: pd.DataFrame,
                                    opening_balances: dict | None = None) -> pd.DataFrame:
    """거래처(+통화) 단위로 기초+당기발생-기말미결제=소멸장부가를 구하고,
    실제결제현금총액과 비교해 실현외환차손익 총액을 역산한다.
    채권(108, 자산성)은 +, 채무(251, 부채성)는 -로 부호를 통일해서 더해야
    양쪽이 섞여도 방향이 맞는다 (채무는 '적게 갚을수록' 이익이라 채권과 반대).
    opening_balances: {(거래처, 통화): 기초장부가(KRW, 부호 통일)} - 없으면 0으로 가정
    (전기 이월 명세서가 없다는 뜻이므로, 결과에 그 가정을 명시한다)."""
    opening_balances = opening_balances or {}

    def _signed_amount(row) -> float:
        amt = row[["차변금액", "대변금액"]].fillna(0).abs().max()
        return amt if row["계정코드"] == AR_ACCOUNT_CODE else -amt

    occur_rows = journal_df[(journal_df["구분"] == "발생") &
                             (journal_df["계정코드"].isin([AR_ACCOUNT_CODE, AP_ACCOUNT_CODE]))].copy()
    occur_rows["_signed"] = occur_rows.apply(_signed_amount, axis=1)
    occur_by_partner = occur_rows.groupby(["거래처", "통화"])["_signed"].sum()

    # 외화예금 라인이 차변(입금=채권 결제)인지 대변(출금=채무 결제)인지로 부호 결정
    cash_rows = journal_df[(journal_df["구분"] == "결제") & (journal_df["계정과목"] == "외화예금")].copy()
    cash_rows["_signed"] = cash_rows.apply(
        lambda r: (r["차변금액"] if pd.notna(r["차변금액"]) else -r["대변금액"]), axis=1
    )
    cash_by_partner = cash_rows.groupby(["거래처", "통화"])["_signed"].sum()

    # 기말 미결제 장부가: 명세서상 기말미결제외화잔액>0인 거래ID의 발생 장부가(부호 포함)를 거래처별로 합산
    outstanding_ids = set(schedule_df.loc[schedule_df["기말미결제외화잔액"] > 0, "전표번호(거래ID)"])
    outstanding_rows = occur_rows[occur_rows["전표번호(거래ID)"].isin(outstanding_ids)]
    outstanding_by_partner = outstanding_rows.groupby(["거래처", "통화"])["_signed"].sum()

    bottom_up_by_partner = screened.groupby(["거래처", "통화"])["회사계상_외환차익차손"].sum()

    all_keys = set(occur_by_partner.index) | set(cash_by_partner.index)
    rows = []
    for key in sorted(all_keys):
        partner, currency = key
        opening = opening_balances.get(key, 0)
        occur_total = occur_by_partner.get(key, 0)
        ending_outstanding = outstanding_by_partner.get(key, 0)
        cash_total = cash_by_partner.get(key, 0)

        consumed_book_value = opening + occur_total - ending_outstanding
        rollforward_gain_loss = cash_total - consumed_book_value
        bottom_up_total = bottom_up_by_partner.get(key, 0)
        diff = round(bottom_up_total - rollforward_gain_loss)

        rows.append({
            "거래처": partner, "통화": currency,
            "기초장부가(가정)": opening, "당기발생장부가(부호포함)": round(occur_total),
            "기말미결제장부가(부호포함)": round(ending_outstanding), "실제결제현금총액(부호포함)": round(cash_total),
            "롤포워드_실현손익": round(rollforward_gain_loss),
            "거래ID매칭_실현손익합계": round(bottom_up_total),
            "차이": diff,
            "판정": "적정(매칭 정합성 확인)" if abs(diff) <= 1000 else "불일치(거래 누락/중복 의심 - 매칭 재확인 필요)",
        })

    return pd.DataFrame(rows)


# ------------------------------------------------------------------
# 5-보조5. 전체 표본 분개 상세 (샘플링된 모든 분개 원본 + 거래별 판정)
# ------------------------------------------------------------------
#
# 지금까지의 시트들은 전부 "거래 단위로 요약된" 뷰라, 실제 분개장 원본 라인이
# 통째로 보이는 시트가 없었다. 검증 대상이 된 분개 전체(=이번에 표본으로
# 들어간 모든 전표)를 원본 라인 그대로 보여주되, 각 거래ID에 대해 이번
# 검증에서 어떤 판정이 나왔는지를 같이 붙여서, "표본 전체 중 어떤 분개가
# 문제였는지"를 한 시트에서 볼 수 있게 한다.

def build_full_journal_detail(journal_df: pd.DataFrame, screened: pd.DataFrame,
                               ye_result: pd.DataFrame) -> pd.DataFrame:
    """검증 대상이 된 분개장 원본 라인 전체에, 거래ID별 최종 판정(결제 스크리닝,
    기말평가, 데이터무결성경고 등)을 합쳐서 보여준다. 문제 있는 거래의 라인이
    맨 위로 오도록 정렬한다."""
    status_map: dict = {}

    for _, r in screened.iterrows():
        txn_id = r["거래ID"]
        settle_date = r["결제일"]
        date_str = settle_date.strftime("%Y-%m-%d") if hasattr(settle_date, "strftime") else str(settle_date)
        status_map.setdefault(txn_id, []).append(f"결제({date_str}): {r['1차플래그']}")
        if pd.notna(r.get("데이터무결성경고")):
            status_map.setdefault(txn_id, []).append(f"⚠ {r['데이터무결성경고']}")

    for _, r in ye_result.iterrows():
        txn_id = r["거래ID"]
        status_map.setdefault(txn_id, []).append(f"기말평가: {r['일치여부']}")

    def _status_for(txn_id):
        return " / ".join(status_map.get(txn_id, ["결제·기말평가 대상 아님(발생만 있거나 검증 범위 밖)"]))

    df = journal_df.copy()
    df["거래_최종판정"] = df["전표번호(거래ID)"].map(_status_for)

    order_key = df["거래_최종판정"].apply(lambda s: {"flag": 0, "neutral": 1, "ok": 2}[_row_status([s])])
    sort_cols = ["전표번호(거래ID)"] + (["라인번호"] if "라인번호" in df.columns else [])
    df = (df.assign(_order=order_key)
            .sort_values(["_order"] + sort_cols, kind="stable")
            .drop(columns="_order"))
    return df


# ------------------------------------------------------------------
# 6. 결과 엑셀 저장
# ------------------------------------------------------------------

FLAG_FILL = PatternFill(start_color="F8CBAD", end_color="F8CBAD", fill_type="solid")
OK_FILL = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
HEADER_FILL_XL = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
HEADER_FONT_XL = Font(bold=True, color="FFFFFF", name="Arial", size=10)
BODY_FONT_XL = Font(name="Arial", size=10)
BOLD_XL = Font(bold=True, name="Arial", size=10)
THIN_BORDER = Border(bottom=Side(style="thin", color="D9D9D9"))

# 어떤 열이든, 셀 값에 이 키워드가 들어있으면 '문제 있음' / '문제 없음'으로 판단.
# (동적으로 날짜 등이 붙는 문구, 예: "기준일 오류 의심 - 2025-04-30..."도 substring으로 잡기 위해
#  정확히 일치하는 값 목록이 아니라 키워드 포함 여부로 판단)
FLAG_KEYWORDS = ["이상치", "부적정", "오류", "불일치", "누락", "미확인", "재검토 필요", "확인 필요"]
OK_KEYWORDS = ["적정", "정상", "통과", "완료"]

# 컬럼명을 사람이 읽기 좋은 표현으로 변경 (내부 로직 컬럼명은 그대로 유지, 엑셀 출력시에만 적용)
COLUMN_LABELS = {
    "회사적용환율(내재)": "회사 적용환율", "회사계상_외환차익차손": "회사 계상액",
    "공식매매기준율": "공식 매매기준율", "공식환율기준일": "환율 기준일",
    "1차플래그": "판정", "재계산_외환차손익(공식환율기준)": "재계산액",
    "회사계상액과의차이(KRW)": "차이(KRW)", "기준일판정": "기준일 판정",
    "추정사용일자": "추정 사용일자", "회사적용환율(기말)": "회사 적용환율(기말)",
    "공식결산환율": "공식 결산환율", "분개장_차변합계": "분개장 차변합계",
    "분개장_대변합계": "분개장 대변합계", "원장_차변합계": "원장 차변합계",
    "원장_대변합계": "원장 대변합계", "회사계상_외환차손익_합계": "회사계상 합계",
    "재계산_외환차손익_합계": "재계산 합계", "차이_합계(KRW)": "차이합계(KRW)",
    "회사평균적용환율": "회사평균환율", "데이터무결성경고": "데이터 무결성 경고",
}

# 열 이름 패턴에 따른 표시 서식
AMOUNT_FORMAT = "#,##0;[RED]-#,##0"
RATE_FORMAT = "#,##0.0000"
PCT_FORMAT = '0.0"%"'
INT_FORMAT = "#,##0"


DATE_FORMAT = "yyyy-mm-dd"


def _column_number_format(col_name: str) -> str | None:
    if "%" in col_name:
        return PCT_FORMAT
    if "환율" in col_name:
        return RATE_FORMAT
    if any(k in col_name for k in ["금액", "차이", "합계", "순액"]):
        return AMOUNT_FORMAT
    if "건수" in col_name:
        return INT_FORMAT
    if col_name.endswith("일") or col_name.endswith("일자"):
        return DATE_FORMAT
    return None


def _row_status(values: list, headers: list = None) -> str:
    """행에 있는 값들을 훑어서 '문제 있음'/'문제 없음'/'중립'을 판단.
    텍스트 판정 컬럼(예: '1차플래그')뿐 아니라, '이상치건수'처럼 숫자로만
    표시되는 컬럼이 0보다 크면 그것도 '문제 있음'으로 인식한다."""
    if headers:
        for h, v in zip(headers, values):
            if "이상치건수" in str(h) and isinstance(v, (int, float)) and v > 0:
                return "flag"
    text = " ".join(str(v) for v in values if v is not None)
    if any(k in text for k in FLAG_KEYWORDS):
        return "flag"
    if any(k in text for k in OK_KEYWORDS):
        return "ok"
    return "neutral"


def _write_table_at(ws, start_row: int, df: pd.DataFrame) -> int:
    """df를 start_row부터 헤더+본문으로 쓰고, 다음에 이어 쓸 수 있는 빈 행 번호를 반환.
    이상치/부적정 행은 빨간색, 적정 행은 초록색으로 통째로 하이라이트하고 맨 위로 정렬,
    차변금액/대변금액이 있으면 순액을 자동 추가, 금액/환율/퍼센트/날짜 열은 서식 자동 적용."""
    if df is None or df.empty:
        ws.cell(row=start_row, column=1, value="(해당 없음)")
        return start_row + 1

    df = df.rename(columns=COLUMN_LABELS)

    if "차변금액" in df.columns and "대변금액" in df.columns:
        debit = pd.to_numeric(df["차변금액"], errors="coerce").fillna(0)
        credit = pd.to_numeric(df["대변금액"], errors="coerce").fillna(0)
        insert_at = df.columns.get_loc("대변금액") + 1
        df = df.copy()
        df.insert(insert_at, "순액", debit - credit)

    status_list = [_row_status(row.tolist(), df.columns.tolist()) for _, row in df.iterrows()]
    order = {"flag": 0, "neutral": 1, "ok": 2}
    df = df.assign(_status=status_list).sort_values(
        "_status", key=lambda s: s.map(order), kind="stable"
    )
    statuses = df["_status"].tolist()
    df = df.drop(columns="_status")

    headers = list(df.columns)
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=start_row, column=c, value=h)
        cell.font = HEADER_FONT_XL
        cell.fill = HEADER_FILL_XL

    number_formats = [_column_number_format(h) for h in headers]
    row_idx = start_row + 1
    for (_, row), status in zip(df.iterrows(), statuses):
        row_fill = FLAG_FILL if status == "flag" else (OK_FILL if status == "ok" else None)
        for c, val in enumerate(row.tolist(), start=1):
            val = None if pd.isna(val) else val
            cell = ws.cell(row=row_idx, column=c, value=val)
            cell.font = BODY_FONT_XL
            cell.border = THIN_BORDER
            if row_fill:
                cell.fill = row_fill
            if number_formats[c - 1] and isinstance(val, (int, float, datetime, pd.Timestamp)):
                cell.number_format = number_formats[c - 1]
        row_idx += 1

    return row_idx


def _autosize_columns(ws):
    """헤더/내용 길이 중 더 긴 쪽 기준으로 열 너비 자동 설정 (최소 10, 최대 34).
    한글 등 CJK 문자는 폭이 영문의 약 1.8배이므로 가중치를 둬서 계산한다."""
    def _display_width(s: str) -> float:
        return sum(1.8 if ord(ch) > 0x2E80 else 1.0 for ch in s)

    for c in range(1, ws.max_column + 1):
        max_len = 0
        for r in range(1, ws.max_row + 1):
            val = ws.cell(row=r, column=c).value
            if val is not None:
                max_len = max(max_len, _display_width(str(val)))
        ws.column_dimensions[get_column_letter(c)].width = min(max(max_len + 2, 10), 34)


def _write_sheet(wb, sheet_name: str, df: pd.DataFrame):
    ws = wb.create_sheet(sheet_name)
    _write_table_at(ws, 1, df)
    _autosize_columns(ws)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def _write_stacked_sheet(wb, sheet_name: str, tables: list):
    """tables: [(제목, df), ...] - 성격이 비슷한 표 여러 개를 한 시트에 세로로 이어붙인다."""
    ws = wb.create_sheet(sheet_name)
    row = 1
    for title, df in tables:
        cell = ws.cell(row=row, column=1, value=title)
        cell.font = Font(bold=True, size=12, name="Arial")
        row += 1
        row = _write_table_at(ws, row, df)
        row += 2  # 표 사이 여백
    _autosize_columns(ws)
    ws.freeze_panes = "A1"


def build_fx_detail_report(screened: pd.DataFrame, ref_check: pd.DataFrame,
                            verified: pd.DataFrame) -> pd.DataFrame:
    """1단계 스크리닝 + 기준일 불일치 탐지 + 2단계 증빙검증을, 거래ID+결제일 기준으로
    하나의 표로 합친다. 원래 3개 시트로 나뉘어 있던 건 사실 같은 거래를 세 각도로
    본 것뿐이라, 하나의 표에서 옆으로 이어 보는 게 오히려 대사하기 편하다."""
    ref_check = ref_check.copy()
    ref_check["결제일"] = pd.to_datetime(ref_check["결제일"])
    merged = screened.merge(
        ref_check[["거래ID", "결제일", "기준일판정", "추정사용일자"]],
        on=["거래ID", "결제일"], how="left",
    )
    if verified is not None and not verified.empty:
        verified = verified.copy()
        verified["결제일"] = pd.to_datetime(verified["결제일"])
        verify_extra = verified[["거래ID", "결제일", "증빙상태", "증빙확인환율", "최종판정", "OCR재확인표본"]]
        merged = merged.merge(verify_extra, on=["거래ID", "결제일"], how="left")
        merged["증빙상태"] = merged["증빙상태"].fillna("증빙검증 대상 아님(1차 스크리닝 적정 통과)")
    return merged


def export_results_to_excel(output_path: str, *, counterparty_summary, rate_trend,
                             screened, agg, ref_check, verified, ye_result, ledger_result,
                             rollforward_result=None, full_detail=None):
    """전체 검증 결과를 시트별로 나눠 하나의 엑셀 워크북으로 저장.
    이상치/부적정 행은 빨간색, 적정 행은 초록색으로 통째로 하이라이트하고 맨 위로 정렬,
    금액/환율/퍼센트 열은 각각 알맞은 서식을 자동 적용한다."""
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    SHEET_GUIDE = [
        ("요약", "합계 중요성(AMPT) 초과 여부와 전체 판정. 제일 먼저 볼 것."),
        ("A.표본전체상세", "검증 대상이 된 분개 원본 라인 전체(발생·결제·기말평가) + 거래별 최종 판정. "
                          "이번에 어떤 분개들을 봤고 그중 뭐가 문제였는지 한 시트에서 확인."),
        ("B.분석적검토", "거래처별 연간 요약 + 월별 환율 추이. 개별 거래를 보기 전에 "
                        "이상한 거래처·시점부터 먼저 짚어내는 분석적 절차용."),
        ("C.외환차손익_상세", "결제 건별 1단계 스크리닝(괴리율) + 기준일 불일치 탐지 + "
                            "2단계 증빙검증을 거래 하나당 한 줄로 합친 상세 표."),
        ("D.외화환산손익", "기말 미결제 건에 대한 재평가 전수 검증(정상/환율오류/재평가누락)."),
        ("E.완전성검증", "계정별원장 대사(분개장 vs 총계정원장) + 거래처별 롤포워드 검증"
                        "(거래ID 매칭 없이 총액으로 재확인)."),
    ]

    ws_guide = wb.create_sheet("안내")
    ws_guide.append(["외환차손익 · 외화환산손익 검증결과 안내"])
    ws_guide["A1"].font = Font(bold=True, size=14, name="Arial")
    ws_guide.append([])
    ws_guide.append(["아래는 각 시트에 대한 설명입니다. 이상치·부적정으로 판정된 행은 빨간색,",])
    ws_guide.append(["적정으로 판정된 행은 초록색으로 표시되며, 문제 있는 행이 항상 위로 정렬됩니다.",])
    ws_guide.append([])
    guide_header_row = ws_guide.max_row + 1
    ws_guide.cell(row=guide_header_row, column=1, value="시트명").font = HEADER_FONT_XL
    ws_guide.cell(row=guide_header_row, column=1).fill = HEADER_FILL_XL
    ws_guide.cell(row=guide_header_row, column=2, value="설명").font = HEADER_FONT_XL
    ws_guide.cell(row=guide_header_row, column=2).fill = HEADER_FILL_XL
    for name, desc in SHEET_GUIDE:
        ws_guide.append([name, desc])
        r = ws_guide.max_row
        ws_guide.cell(row=r, column=1).font = BOLD_XL
        ws_guide.cell(row=r, column=2).font = BODY_FONT_XL
        ws_guide.cell(row=r, column=1).border = THIN_BORDER
        ws_guide.cell(row=r, column=2).border = THIN_BORDER
    ws_guide.append([])
    r = ws_guide.max_row + 1
    ws_guide.cell(row=r, column=1, value="색상 범례").font = BOLD_XL
    r += 1
    ws_guide.cell(row=r, column=1, value="이상치 / 부적정 (확인 필요)").fill = FLAG_FILL
    ws_guide.cell(row=r, column=1).font = BODY_FONT_XL
    r += 1
    ws_guide.cell(row=r, column=1, value="적정 (통과)").fill = OK_FILL
    ws_guide.cell(row=r, column=1).font = BODY_FONT_XL
    ws_guide.column_dimensions["A"].width = 22
    ws_guide.column_dimensions["B"].width = 70

    ws0 = wb.create_sheet("요약")
    ws0.append(["외환차손익·외화환산손익 검증 결과 요약"])
    ws0["A1"].font = Font(bold=True, size=14, name="Arial")
    ws0.append([])
    summary_rows = [
        ("합계 중요성(AMPT)", agg["AMPT"], AMOUNT_FORMAT),
        ("순차이합계(KRW)", agg["순차이합계(KRW)"], AMOUNT_FORMAT),
        ("절대값차이합계(KRW)", agg["절대값차이합계(KRW)"], AMOUNT_FORMAT),
        ("판정", agg["판정"], None),
        ("확인불가건수(환율조회실패 등)", agg["확인불가건수"], INT_FORMAT),
    ]
    for label, value, fmt in summary_rows:
        ws0.append([label, value])
        r = ws0.max_row
        ws0.cell(row=r, column=1).font = BOLD_XL
        ws0.cell(row=r, column=2).font = BODY_FONT_XL
        if fmt:
            ws0.cell(row=r, column=2).number_format = fmt
        if _row_status([label, value]) == "flag":
            ws0.cell(row=r, column=2).fill = FLAG_FILL
    ws0.column_dimensions["A"].width = 22
    ws0.column_dimensions["B"].width = 40

    _write_sheet(wb, "A.표본전체상세", full_detail)
    _write_stacked_sheet(wb, "B.분석적검토", [
        ("거래처별 연간 요약", counterparty_summary),
        ("월별 환율 추이", rate_trend),
    ])
    fx_detail = build_fx_detail_report(screened, ref_check, verified)
    _write_sheet(wb, "C.외환차손익_상세", fx_detail)
    _write_sheet(wb, "D.외화환산손익", ye_result)
    completeness_tables = [("계정별원장 대사", ledger_result)]
    if rollforward_result is not None:
        completeness_tables.append(("거래처별 롤포워드 검증", rollforward_result))
    _write_stacked_sheet(wb, "E.완전성검증", completeness_tables)

    wb.save(output_path)
    print(f"\n결과 엑셀 저장 완료: {output_path}")


# ------------------------------------------------------------------
# 6-보조. 감사조서(Word) 저장
# ------------------------------------------------------------------

def _first_non_null(row, cols: list):
    """row(Series)에서 cols 순서대로 값이 있는(NaN/None이 아닌) 첫 값을 반환.
    여러 판정 컬럼(1차플래그/기준일판정/최종판정 등) 중 있는 것만 이어붙이는 데 사용."""
    for c in cols:
        v = row.get(c) if hasattr(row, "get") else None
        if v is not None and not pd.isna(v):
            return v
    return None


def _collect_exception_rows(구분: str, df: pd.DataFrame, id_col: str, judge_cols: list,
                             diff_col: str = None) -> list:
    """df를 훑어 _row_status()가 'flag'로 판정한 행만 감사조서 Exception 표 형식으로 뽑아낸다.
    새 판정을 만들지 않고, Excel 하이라이트에 쓰는 것과 동일한 _row_status()/FLAG_KEYWORDS를
    그대로 재사용해 Excel과 Word 산출물의 Exception 목록이 항상 일치하도록 한다."""
    if df is None or df.empty:
        return []
    rows = []
    headers = df.columns.tolist()
    for _, row in df.iterrows():
        if _row_status(row.tolist(), headers) != "flag":
            continue
        judge_parts = []
        for c in judge_cols:
            v = row.get(c)
            if v is not None and not pd.isna(v) and "정상" not in str(v) and "대상 아님" not in str(v):
                judge_parts.append(str(v))
        diff = row.get(diff_col) if diff_col else None
        rows.append({
            "구분": 구분,
            "식별자": row.get(id_col),
            "통화": row.get("통화", "-"),
            "판정": " / ".join(judge_parts) if judge_parts else "부적정",
            "차이금액(KRW)": diff if diff is not None and not pd.isna(diff) else None,
        })
    return rows


def generate_audit_workpaper(output_path: str, *, screened: pd.DataFrame, agg: dict,
                              ref_check: pd.DataFrame, verified: pd.DataFrame,
                              ye_result: pd.DataFrame, ledger_result: pd.DataFrame,
                              rollforward_result: pd.DataFrame, year_end_date: str) -> None:
    """검증결과.xlsx와 별도로, 이미 계산된 결과(screened/agg/ye_result/ledger_result/
    rollforward_result 등)를 감사조서 표준 형식(검증 목적 - 검증 대상 - 수행 절차 -
    검증 결과 - 결론)의 Word 문서로 옮겨 적는다. 여기서 새로운 판단(중요성 판단, 개별 건의
    적정/부적정 판정)을 내리지 않고, 파이프라인이 이미 만든 판정 문자열과 플래그
    (_row_status/FLAG_KEYWORDS - Excel 하이라이트와 동일 기준)를 그대로 인용만 한다."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(10)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    doc.add_heading("감사조서 - 외환차손익 및 외화환산손익 검증", level=0)
    doc.add_paragraph(f"결산일: {year_end_date}    작성일시(자동생성): "
                       f"{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph().add_run(
        "본 조서는 검증 파이프라인의 산출물을 자동으로 정리한 것으로, "
        "최종 감사의견 및 중요성 판단은 감사인이 수행한다."
    ).italic = True

    # 1. 검증 목적
    doc.add_heading("1. 검증 목적", level=1)
    doc.add_paragraph(
        f"결산일({year_end_date}) 현재 외화화폐성 항목에 대해 계상된 외환차손익(실현손익) "
        "및 외화환산손익(미실현손익)이 적용환율 및 관련 회계처리 기준에 따라 적정하게 "
        "계상되었는지 확인한다."
    )

    # 2. 검증 대상
    doc.add_heading("2. 검증 대상", level=1)
    counterparty_cnt = screened["거래처"].nunique() if "거래처" in screened.columns else 0
    doc.add_paragraph(f"대상 거래처 수(결제 건 기준): {counterparty_cnt}개")
    settle_by_cur = (screened.groupby("통화").agg(결제건수=("거래ID", "count"),
                                                  결제외화금액합계=("외화금액", "sum"))
                      if not screened.empty else pd.DataFrame(columns=["결제건수", "결제외화금액합계"]))
    ye_by_cur = (ye_result.groupby("통화").agg(기말미결제건수=("거래ID", "count"),
                                               기말외화금액합계=("외화금액", "sum"))
                 if not ye_result.empty else pd.DataFrame(columns=["기말미결제건수", "기말외화금액합계"]))
    currency_summary = settle_by_cur.join(ye_by_cur, how="outer").fillna(0).reset_index()
    currency_summary = currency_summary.rename(columns={"index": "통화"})
    if currency_summary.empty:
        doc.add_paragraph("(해당 없음)")
    else:
        table = doc.add_table(rows=1, cols=len(currency_summary.columns))
        table.style = "Table Grid"
        for c, h in enumerate(currency_summary.columns):
            table.rows[0].cells[c].text = str(h)
        for _, row in currency_summary.iterrows():
            cells = table.add_row().cells
            for c, val in enumerate(row):
                cells[c].text = f"{val:,.0f}" if isinstance(val, (int, float)) else str(val)

    # 3. 수행 절차
    doc.add_heading("3. 수행 절차", level=1)
    procedures = [
        "1단계 스크리닝: 결제 건별로 분개상 원화금액·외화금액에서 내재환율을 역산해, "
        "결제일의 한국수출입은행 공식 매매기준율과 비교한다. 괴리율이 5%를 초과하는 "
        "거래를 이상치로 플래그한다.",
        "합계 중요성(AMPT) 검증: 개별 건이 괴리율 기준을 통과하더라도, 재계산액과 "
        "회사계상액의 차이를 전체 합산해 허용가능 오류금액(AMPT) 초과 여부를 확인한다.",
        "결제일-환율기준일 불일치 탐지: 회사 적용환율이 결제일이 아닌 다른 날짜(전월말, "
        "전영업일 등)의 공식 환율과 정확히 일치하는지 확인해 기준일 오적용을 탐지한다.",
        "2단계 증빙 대조(Vision OCR): 1단계 이상치로 플래그된 거래에 대해 은행 "
        "외환거래확인서 등 증빙을 Claude Vision으로 읽어 실제 적용환율을 추출하고 "
        "회사 계상액과 최종 대사한다.",
        "기말환산 검증: 결산일 현재 미결제 외화자산·부채를 결산일 공식 환율로 전수 "
        "재평가해 기말 환산 처리의 적정성을 확인한다(재평가 누락 건 포함).",
        "계정별원장 대사: 분개장에서 집계한 외환차익·차손·외화환산이익·손실 계정 합계가 "
        "총계정원장 잔액과 일치하는지 확인해 수기 조정 등으로 인한 완전성 이슈를 탐지한다.",
        "롤포워드 검증: 발생-결제 거래를 1:1 매칭하기 어려운 경우에도, 거래처·통화 "
        "단위로 기초+당기발생-기말잔액=결제액의 원리로 총액 정합성을 재확인한다.",
    ]
    for i, text in enumerate(procedures, start=1):
        doc.add_paragraph(f"{i}) {text}", style="List Number")

    # 4. 검증 결과
    doc.add_heading("4. 검증 결과", level=1)
    total_population = len(screened) + len(ye_result)
    doc.add_paragraph(f"전체 검증 대상: {total_population:,}건 "
                       f"(외환차손익 결제 건 {len(screened):,}건 + 외화환산손익 기말 미결제 건 "
                       f"{len(ye_result):,}건)")
    doc.add_paragraph(f"합계 중요성(AMPT): {agg['AMPT']:,.0f}원 / "
                       f"절대값차이합계(KRW): {agg['절대값차이합계(KRW)']:,.0f}원 / "
                       f"판정: {agg['판정']}")
    if agg.get("확인불가건수", 0):
        doc.add_paragraph(f"환율조회 실패 등으로 확인 불가능한 건: {agg['확인불가건수']:,}건")

    fx_detail = build_fx_detail_report(screened, ref_check, verified)
    exception_rows = []
    exception_rows += _collect_exception_rows(
        "외환차손익(결제)", fx_detail, "거래ID",
        ["1차플래그", "기준일판정", "최종판정"], "회사계상액과의차이(KRW)")
    exception_rows += _collect_exception_rows(
        "외화환산손익(기말)", ye_result, "거래ID", ["일치여부"])
    exception_rows += _collect_exception_rows(
        "완전성검증(원장대사)", ledger_result, "계정과목", ["일치여부"])
    exception_rows += _collect_exception_rows(
        "완전성검증(롤포워드)", rollforward_result, "거래처", ["판정"], "차이")

    doc.add_paragraph(f"Exception(이상치/부적정) 건수: {len(exception_rows):,}건")
    if exception_rows:
        cols = ["구분", "식별자", "통화", "판정", "차이금액(KRW)"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for c, h in enumerate(cols):
            table.rows[0].cells[c].text = h
        for exc in exception_rows:
            cells = table.add_row().cells
            for c, key in enumerate(cols):
                val = exc[key]
                if key == "차이금액(KRW)" and isinstance(val, (int, float)):
                    cells[c].text = f"{val:,.0f}"
                else:
                    cells[c].text = str(val) if val is not None else "-"
    else:
        doc.add_paragraph("(해당 없음)")

    if "OCR재확인표본" in fx_detail.columns:
        recheck_count = fx_detail["OCR재확인표본"].notna().sum()
        doc.add_paragraph(
            f"OCR 재확인 표본(QC): {recheck_count:,}건 — OCR이 '증빙과 일치'로 판정한 건 중 "
            "위음성 여부를 감사인이 별도로 재확인하기 위해 자동 추출한 표본이다. Exception과는 "
            "별개이며, 재확인 결과 OCR 오류가 발견되면 표본 확대 또는 전수 재검토로 전환한다."
        )

    # 5. 결론
    doc.add_heading("5. 결론", level=1)
    breach = agg["중요성초과여부"]
    if breach:
        conclusion = (f"합계 중요성 검증 결과, 절대값차이합계 {agg['절대값차이합계(KRW)']:,.0f}원이 "
                      f"허용가능 오류금액(AMPT) {agg['AMPT']:,.0f}원을 초과하여 중요한 차이가 "
                      "발견되었다.")
    else:
        conclusion = (f"합계 중요성 검증 결과, 절대값차이합계 {agg['절대값차이합계(KRW)']:,.0f}원이 "
                      f"허용가능 오류금액(AMPT) {agg['AMPT']:,.0f}원 이내로, 중요한 차이는 발견되지 "
                      "않았다.")
    if exception_rows:
        conclusion += f" 상기 Exception {len(exception_rows):,}건에 대해서는 추가 검토가 필요하다."
    else:
        conclusion += " Exception으로 식별된 건은 없다."
    doc.add_paragraph(conclusion)

    doc.save(output_path)
    print(f"감사조서 저장 완료: {output_path}")


# ------------------------------------------------------------------
# 7. 실행
# ------------------------------------------------------------------

def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="외환차손익/외화환산손익 검증 파이프라인. 인자 없이 실행하면 sample_data/의 "
                     "샘플 분개장·명세서·계정별원장을 대상으로 동작합니다."
    )
    parser.add_argument("--journal", default="sample_data/분개장.xlsx", help="분개장 엑셀 경로")
    parser.add_argument("--schedule", default="sample_data/명세서_외화자산부채명세서.xlsx",
                         help="외화자산부채명세서 엑셀 경로")
    parser.add_argument("--ledger", default="sample_data/계정별원장.xlsx", help="계정별원장 엑셀 경로")
    parser.add_argument("--year-end-date", default="2025-12-31", help="결산일 (YYYY-MM-DD)")
    parser.add_argument("--output", default="검증결과.xlsx", help="결과 엑셀 저장 경로")
    parser.add_argument("--ampt", type=float, default=AMPT,
                         help="허용가능 오류금액(Tolerable Misstatement). 기본값은 FX_AMPT 환경변수")
    parser.add_argument("--workpaper", action="store_true",
                         help="검증결과.xlsx와 함께 감사조서(.docx)도 생성")
    parser.add_argument("--workpaper-output", default="감사조서_외환차손익.docx",
                         help="감사조서 Word 파일 저장 경로 (--workpaper와 함께 사용)")
    return parser


def main(argv=None):
    args = build_arg_parser().parse_args(argv)

    journal = load_journal(args.journal)
    schedule = pd.read_excel(args.schedule)
    ledger = pd.read_excel(args.ledger)

    # 0단계: 실무 감사 절차와 동일하게, 거래 하나하나를 보기 전에
    # 거래처별 연간 요약 + 월별 환율 추이부터 확인한다.
    settlements = extract_settlement_transactions(journal)
    screened_preview = screen_fx_settlements(settlements)  # 요약 산출을 위해 먼저 1회 계산

    print("=== 0단계. 거래처별 연간 요약 (분석적 절차 - 여기서 이상해 보이는 거래처만 드릴다운) ===")
    counterparty_summary = build_counterparty_summary(screened_preview)
    print(counterparty_summary.to_string(index=False))

    print("\n=== 0단계-보조. 월별 환율 추이 (연중 이상 시점 파악용) ===")
    rate_trend = build_rate_trend_summary(screened_preview)
    print(rate_trend.to_string(index=False))

    print("\n=== A. 외환차손익 1단계 스크리닝 (거래 단위 상세) ===")
    if settlements["데이터무결성경고"].notna().any():
        print("⚠ 데이터 무결성 경고 발견:")
        print(settlements.loc[settlements["데이터무결성경고"].notna(), ["거래ID", "데이터무결성경고"]].to_string(index=False))
    screened = screened_preview
    print(screened[["거래ID", "통화", "회사적용환율(내재)", "공식매매기준율", "괴리율(%)",
                     "1차플래그", "회사계상액과의차이(KRW)"]])

    print("\n=== A-보조. 합계 중요성 검증 (개별 통과 건 포함 전체 합산) ===")
    agg = check_aggregate_materiality(screened, args.ampt)
    print(agg)

    print("\n=== A-보조2. 결제일-환율기준일 불일치 탐지 (5% 통과 건 포함 전수 검사) ===")
    ref_check = detect_reference_date_mismatch(settlements)
    print(ref_check[["거래ID", "결제일", "기준일판정", "추정사용일자"]])

    print("\n=== A. 외환차손익 2단계 증빙검증 ===")
    if agg["중요성초과여부"]:
        print("↳ 합계 중요성 초과 - 개별 통과 건까지 포함해 전체를 정밀검증 대상으로 확장합니다.")
        to_verify = screened
    else:
        to_verify = screened[screened["1차플래그"] == "이상치(정밀검증필요)"]
    verified = verify_with_evidence(to_verify)
    print(verified[["거래ID", "증빙상태", "증빙확인환율", "최종판정"]])

    print("\n=== B. 외화환산손익 전수 검증 ===")
    ye_txns = extract_yearend_transactions(journal)
    missing_ye = get_unsettled_yearend_candidates(journal, schedule)
    ye_result = verify_yearend_translation(ye_txns, missing_ye, args.year_end_date)
    print(ye_result[["거래ID", "통화", "일치여부"]])

    print("\n=== C. 계정별원장 대사 (분개장 <-> 총계정원장 교차검증) ===")
    ledger_result = verify_ledger_reconciliation(journal, ledger)
    print(ledger_result)

    print("\n=== D. 거래처별 롤포워드 검증 (거래ID 매칭 없이 총액 정합성 확인) ===")
    rollforward_result = build_rollforward_verification(journal, schedule, screened)
    print(rollforward_result)

    full_detail = build_full_journal_detail(journal, screened, ye_result)

    export_results_to_excel(
        args.output,
        counterparty_summary=counterparty_summary, rate_trend=rate_trend,
        screened=screened, agg=agg, ref_check=ref_check, verified=verified,
        ye_result=ye_result, ledger_result=ledger_result, rollforward_result=rollforward_result,
        full_detail=full_detail,
    )

    if args.workpaper:
        generate_audit_workpaper(
            args.workpaper_output,
            screened=screened, agg=agg, ref_check=ref_check, verified=verified,
            ye_result=ye_result, ledger_result=ledger_result,
            rollforward_result=rollforward_result, year_end_date=args.year_end_date,
        )


if __name__ == "__main__":
    main()
